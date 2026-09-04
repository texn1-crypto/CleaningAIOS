from __future__ import annotations

import hashlib
import os
import re
from datetime import datetime, timezone
from html import escape
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy import select
from sqlalchemy.orm import Session

from .config import settings
from .models import BusinessRecord
from .proposals import _register_font, _safe_filename


ALLOWED_SOURCE_SUFFIXES = {".docx", ".pdf"}
NAVY = "173F5F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "1F7A70"
MUTED = "66737F"
LIGHT = "F4F6F9"
INK = "1C2933"


def _storage_path(value: str) -> Path:
    root = Path(settings.document_storage_path).resolve()
    path = Path(value).resolve()
    try:
        path.relative_to(root)
    except ValueError as exc:
        raise ValueError("Исходный документ находится вне защищённого хранилища.") from exc
    if not path.is_file() or path.suffix.lower() not in ALLOWED_SOURCE_SUFFIXES:
        raise ValueError("Исходный DOCX/PDF недоступен в защищённом хранилище.")
    if path.stat().st_size > settings.max_document_bytes:
        raise ValueError("Исходный документ превышает разрешённый размер.")
    return path


def _extract_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        document = Document(path)
        values = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    values.append(" | ".join(cells))
        text = "\n".join(values)
    else:
        reader = PdfReader(str(path))
        text = "\n".join((page.extract_text() or "").strip() for page in reader.pages[:100])
    return text[:200_000]


def _client_name(filename: str, text: str) -> str:
    candidate = Path(filename).stem
    match = re.search(r"\b(?:ЖК|БЦ|УК|ТСЖ)\s+[\w\- «»\"Ёё]{2,80}", candidate, flags=re.IGNORECASE)
    if match:
        return " ".join(match.group(0).strip(" -_").split())
    candidate = re.sub(r"^\s*кп\s+", "", candidate, flags=re.IGNORECASE)
    candidate = re.sub(r"\bип\s+соколов\b.*?(?=\b(?:жк|бц|ук|тсж)\b|$)", "", candidate, flags=re.IGNORECASE)
    cleaned = " ".join(candidate.strip(" -_").split())
    return cleaned[:120] or "Заказчика"


def _fact_lines(text: str) -> list[str]:
    compact = " ".join(text.split())
    structured_patterns = (
        ("Уборка МОП", r"Уборка\s+МОП\s*\(([^)]*уборщиц[^)]*)\)\s*(\d{1,3}(?:\s+\d{3})+)", True),
        ("Снабжение", r"Снабжение\s+(\d{1,3}(?:\s+\d{3})+)", False),
        ("Придомовая территория", r"Придомовая\s+территория\s*\(([^)]*дворник[^)]*)\)\s*(\d{1,3}(?:\s+\d{3})+)", True),
        ("Менеджер клининга", r"Менеджер\s+клининга\s+(\d{1,3}(?:\s+\d{3})+?)(?=\s+\d{1,3}(?:\s+\d{3})+\s+руб)", False),
    )
    structured: list[str] = []
    for label, pattern, has_scope in structured_patterns:
        match = re.search(pattern, compact, flags=re.IGNORECASE)
        if not match:
            continue
        scope_value = " ".join(match.group(1).split()) if has_scope else ""
        scope_value = re.sub(r"\b2\s+уборщица\b", "2 уборщицы", scope_value, flags=re.IGNORECASE)
        scope = f" ({scope_value})" if has_scope else ""
        amount = " ".join(match.group(2 if has_scope else 1).split())
        structured.append(f"{label}{scope} — {amount} руб.")
    total = re.search(
        r"Менеджер\s+клининга\s+\d{1,3}(?:\s+\d{3})+\s+(\d{1,3}(?:\s+\d{3})+)\s+руб",
        compact,
        flags=re.IGNORECASE,
    )
    if total:
        structured.append(f"Итоговая стоимость — {' '.join(total.group(1).split())} руб.")
    if len(structured) >= 3:
        return structured

    keywords = (
        "руб", "цен", "стоим", "м²", "м2", "кв.", "график", "смен", "час", "раз в", "ежедн", "адрес",
        "площад", "период", "месяц", "ндс", "оплат",
    )
    seen: set[str] = set()
    facts: list[str] = []
    for raw in text.splitlines():
        line = " ".join(raw.split()).strip(" |•-")
        normalized = line.lower().replace("ё", "е")
        if len(line) < 5 or len(line) > 240 or normalized in seen:
            continue
        if any(word in normalized for word in keywords) and (any(char.isdigit() for char in line) or "адрес" in normalized):
            seen.add(normalized)
            facts.append(line)
        if len(facts) >= 12:
            break
    return facts


def build_professional_copy(payload: dict[str, Any]) -> dict[str, Any]:
    path = _storage_path(str(payload.get("source_path", "")))
    source_text = _extract_text(path)
    if not source_text.strip():
        raise ValueError("В исходном документе не найден текст для профессиональной редакции.")
    source_filename = str(payload.get("source_filename") or path.name)
    client = _client_name(source_filename, source_text)
    facts = _fact_lines(source_text)
    source_email = re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-zА-Яа-я]{2,}", source_text)
    source_phone = re.search(r"(?:\+7|8)[\s()\-\d]{9,20}\d", source_text)
    sections = {
        "opening": (
            f"Уважаемые коллеги! {settings.company_name} предлагает комплексное решение по профессиональному "
            f"клинингу для {client}. Наша задача — обеспечить стабильный результат, прозрачный контроль качества "
            "и предсказуемую работу команды на объекте."
        ),
        "value": (
            "Мы проектируем услугу от фактической нагрузки и сценариев использования зон. После осмотра объекта фиксируем зоны, "
            "периодичность, состав смен, технологические карты и критерии приёмки. Это позволяет управлять не процессом «уборки», а измеримым уровнем чистоты."
        ),
        "quality": (
            "За объектом закрепляется ответственный. Качество проверяется по чек-листам и контрольным точкам; замечания регистрируются, назначаются исполнителю "
            "и контролируются до устранения. Формат отчётности и SLA согласуются с заказчиком до запуска."
        ),
        "launch": (
            "Запуск проходит поэтапно: осмотр и уточнение объёма, согласование регламента и графика, подбор команды и инвентаря, стартовый контроль и корректировка технологических карт. "
            "Финальные цена, график и объём обязательств закрепляются только в согласованном договоре."
        ),
    }
    checksum = hashlib.sha256(path.read_bytes()).hexdigest()
    return {
        "status": "ready",
        "provider": "private_local_copy_engine",
        "external_ai_used": False,
        "client_name": client,
        "sections": sections,
        "source_facts": facts,
        "source_contacts": {
            "email": source_email.group(0) if source_email else "",
            "phone": " ".join(source_phone.group(0).split()) if source_phone else "",
        },
        "source_text_characters": len(source_text),
        "source_checksum_sha256": checksum,
        "privacy": "source_document_not_sent_to_external_ai",
        "evidence": [{"type": "proposal_source_text", "checksum_sha256": checksum, "characters": len(source_text)}],
    }


def build_creative_direction(payload: dict[str, Any]) -> dict[str, Any]:
    copy = payload.get("copy") if isinstance(payload.get("copy"), dict) else {}
    if copy.get("status") != "ready" or not copy.get("sections"):
        raise ValueError("Текстовый агент не подготовил проверяемую редакцию.")
    return {
        "status": "ready",
        "provider": "private_local_design_engine",
        "preset": "narrative_proposal",
        "header_pattern": "proposal_centerpiece",
        "page_size": "Letter",
        "margins_inches": 1.0,
        "body": {"font": "Calibri", "size_pt": 11, "after_pt": 8, "line_spacing": 1.333},
        "palette": {"navy": f"#{NAVY}", "blue": f"#{BLUE}", "teal": f"#{TEAL}", "light": f"#{LIGHT}"},
        "owner_approval_watermark": True,
        "evidence": [{"type": "proposal_design_spec", "preset": "narrative_proposal", "header_pattern": "proposal_centerpiece"}],
    }


def _set_font(run, *, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def _shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _table_geometry(table, widths: list[int]) -> None:
    table.style = "Table Grid"
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_width.set(qn("w:w"), str(widths[index]))
            tc_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _configure_docx_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, DARK_BLUE),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    _set_font(header.add_run(f"{settings.company_name}  |  Коммерческое предложение"), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_font(footer.add_run("Проект для утверждения владельцем  |  "), size=8, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def _add_body(document: Document, copy: dict[str, Any]) -> None:
    client = str(copy["client_name"])
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(8)
    _set_font(kicker.add_run(settings.company_name.upper()), size=11, color=TEAL, bold=True)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    _set_font(title.add_run("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ"), size=24, color=NAVY, bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    _set_font(subtitle.add_run(f"Комплексный клининг для {client}"), size=14, color=MUTED)

    badge = document.add_table(rows=1, cols=1)
    _table_geometry(badge, [9360])
    _shade(badge.cell(0, 0), NAVY)
    p = badge.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    _set_font(p.add_run("ПРОЕКТ · ТРЕБУЕТ ПРОВЕРКИ И УТВЕРЖДЕНИЯ ВЛАДЕЛЬЦА"), size=9, color="FFFFFF", bold=True)
    document.add_paragraph()

    sections = copy["sections"]
    headings = (
        ("Предложение", "opening"),
        ("Решение под задачи объекта", "value"),
        ("Управление качеством", "quality"),
        ("Порядок запуска", "launch"),
    )
    for heading, key in headings:
        document.add_paragraph(heading, style="Heading 1")
        paragraph = document.add_paragraph(sections[key])
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    document.add_page_break()
    document.add_paragraph("Коммерческие условия из исходного КП", style="Heading 1")
    facts = copy.get("source_facts") or []
    if facts:
        table = document.add_table(rows=1, cols=2)
        _table_geometry(table, [2700, 6660])
        for index, label in enumerate(("Позиция", "Данные из исходника")):
            _shade(table.cell(0, index), LIGHT)
            paragraph = table.cell(0, index).paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            _set_font(paragraph.add_run(label), size=9.5, color=NAVY, bold=True)
        for index, fact in enumerate(facts, start=1):
            cells = table.add_row().cells
            for cell in cells:
                _set_cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            left = cells[0].paragraphs[0]
            left.paragraph_format.space_after = Pt(0)
            _set_font(left.add_run(f"Условие {index}"), size=9.5, color=MUTED, bold=True)
            right = cells[1].paragraphs[0]
            right.paragraph_format.space_after = Pt(0)
            _set_font(right.add_run(str(fact)), size=9.5, color=INK)
        _table_geometry(table, [2700, 6660])
        note = document.add_paragraph()
        note.paragraph_format.space_before = Pt(4)
        note.paragraph_format.space_after = Pt(4)
        _set_font(note.add_run("Все перенесённые условия требуют проверки владельца перед отправкой заказчику."), size=9, color=MUTED)
    else:
        document.add_paragraph("Числовые условия не переносились автоматически. Цена, график и объём нужно утвердить после аудита объекта.")

    document.add_paragraph("Следующий шаг", style="Heading 1")
    document.add_paragraph("Предлагаем провести осмотр объекта и короткую рабочую встречу. По итогам мы подготовим финальный регламент, расчёт ресурсов и стоимость.")
    closing = document.add_table(rows=1, cols=1)
    _table_geometry(closing, [9360])
    _shade(closing.cell(0, 0), LIGHT)
    paragraph = closing.cell(0, 0).paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    source_contacts = copy.get("source_contacts") or {}
    contact = source_contacts.get("phone") or source_contacts.get("email") or settings.company_phone or settings.company_email or "требует заполнения"
    _set_font(paragraph.add_run(
        f"{settings.company_legal_name or settings.company_name}\n"
        f"ИНН: {settings.company_inn or 'требует заполнения'}\n"
        f"Контакт из исходного КП: {contact} (проверить актуальность)"
    ), size=9.5, color=INK)


def _build_docx(path: Path, copy: dict[str, Any]) -> None:
    document = Document()
    _configure_docx_styles(document)
    document.core_properties.title = f"Коммерческое предложение для {copy['client_name']}"
    document.core_properties.author = settings.company_legal_name or settings.company_name
    document.core_properties.subject = "Проект для утверждения владельцем"
    _add_body(document, copy)
    document.save(path)


def _build_pdf(path: Path, copy: dict[str, Any]) -> None:
    font = _register_font()
    styles = getSampleStyleSheet()
    title = ParagraphStyle("StudioTitle", parent=styles["Title"], fontName=font, fontSize=23, leading=28, textColor=colors.HexColor(f"#{NAVY}"), alignment=TA_CENTER, spaceAfter=8)
    subtitle = ParagraphStyle("StudioSubtitle", parent=styles["BodyText"], fontName=font, fontSize=13, leading=17, textColor=colors.HexColor(f"#{MUTED}"), alignment=TA_CENTER, spaceAfter=22)
    h1 = ParagraphStyle("StudioH1", parent=styles["Heading1"], fontName=font, fontSize=16, leading=20, textColor=colors.HexColor(f"#{BLUE}"), spaceBefore=18, spaceAfter=10)
    body = ParagraphStyle("StudioBody", parent=styles["BodyText"], fontName=font, fontSize=10.5, leading=16, textColor=colors.HexColor(f"#{INK}"), alignment=TA_JUSTIFY, spaceAfter=8)
    small = ParagraphStyle("StudioSmall", parent=body, fontSize=8.5, leading=12, textColor=colors.HexColor(f"#{MUTED}"), alignment=TA_LEFT)
    badge = ParagraphStyle("StudioBadge", parent=small, fontSize=8.5, leading=11, textColor=colors.white, alignment=TA_CENTER)
    document = SimpleDocTemplate(str(path), pagesize=LETTER, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch, title=f"КП для {copy['client_name']}", author=settings.company_legal_name or settings.company_name)

    def pdf_escape(value: Any) -> str:
        return escape(str(value)).replace("—", "-").replace("–", "-")

    def page(canvas, doc):
        canvas.saveState()
        canvas.setFont(font, 8)
        canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
        canvas.drawString(inch, 0.48 * inch, f"{settings.company_name}  |  проект для утверждения")
        canvas.drawRightString(LETTER[0] - inch, 0.48 * inch, f"Страница {doc.page}")
        canvas.restoreState()

    story: list[Any] = [
        Paragraph(pdf_escape(settings.company_name.upper()), ParagraphStyle("Kicker", parent=small, textColor=colors.HexColor(f"#{TEAL}"), fontSize=10.5, alignment=TA_CENTER, spaceAfter=8)),
        Paragraph("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ", title),
        Paragraph(f"Комплексный клининг для {pdf_escape(copy['client_name'])}", subtitle),
        Table([[Paragraph("ПРОЕКТ · ТРЕБУЕТ ПРОВЕРКИ И УТВЕРЖДЕНИЯ ВЛАДЕЛЬЦА", badge)]], colWidths=[6.5 * inch], style=TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{NAVY}")), ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8)])),
        Spacer(1, 10),
    ]
    for heading, key in (("Предложение", "opening"), ("Решение под задачи объекта", "value"), ("Управление качеством", "quality"), ("Порядок запуска", "launch")):
        story.extend([Paragraph(heading, h1), Paragraph(pdf_escape(copy["sections"][key]), body)])
    story.extend([PageBreak(), Paragraph("Коммерческие условия из исходного КП", h1)])
    facts = copy.get("source_facts") or []
    if facts:
        rows = [[Paragraph("Позиция", small), Paragraph("Данные из исходника", small)]]
        rows.extend([[Paragraph(f"Условие {index}", small), Paragraph(pdf_escape(fact), small)] for index, fact in enumerate(facts, start=1)])
        story.append(Table(rows, colWidths=[1.875 * inch, 4.625 * inch], repeatRows=1, style=TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{LIGHT}")), ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D5DCE1")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8), ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 7)])))
        story.append(Paragraph("Все перенесённые условия требуют проверки владельца перед отправкой заказчику.", small))
    else:
        story.append(Paragraph("Цена, график и объём нужно утвердить после аудита объекта.", body))
    source_contacts = copy.get("source_contacts") or {}
    contact = source_contacts.get("phone") or source_contacts.get("email") or settings.company_phone or settings.company_email or "требует заполнения"
    story.extend([
        Paragraph("Следующий шаг", h1),
        Paragraph("Предлагаем провести осмотр объекта и короткую рабочую встречу. По итогам мы подготовим финальный регламент, расчёт ресурсов и стоимость.", body),
        Spacer(1, 12),
        Table([[Paragraph(
            f"{escape(settings.company_legal_name or settings.company_name)}<br/>"
            f"ИНН: {escape(settings.company_inn or 'требует заполнения')}<br/>"
            f"Контакт из исходного КП: {escape(contact)} (проверить актуальность)",
            small,
        )]], colWidths=[6.5 * inch], style=TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{LIGHT}")), ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#D5DCE1")), ("LEFTPADDING", (0, 0), (-1, -1), 10), ("RIGHTPADDING", (0, 0), (-1, -1), 10), ("TOPPADDING", (0, 0), (-1, -1), 9), ("BOTTOMPADDING", (0, 0), (-1, -1), 9)])),
    ])
    document.build(story, onFirstPage=page, onLaterPages=page)


def create_proposal_revision(db: Session, payload: dict[str, Any], copy: dict[str, Any], design: dict[str, Any]) -> dict[str, Any]:
    source = _storage_path(str(payload.get("source_path", "")))
    request_text = str(payload.get("request_text", ""))[:4000]
    signature = hashlib.sha256((copy["source_checksum_sha256"] + request_text).encode()).hexdigest()
    external_id = f"KPR-{signature[:20]}"
    existing = db.scalar(select(BusinessRecord).where(BusinessRecord.record_type == "proposal_revision", BusinessRecord.external_id == external_id))
    if existing:
        files = existing.data.get("files") or {}
        if all(Path(str(item.get("storage_path", ""))).is_file() for item in files.values() if isinstance(item, dict)):
            return _revision_result(existing, copy_task_id=payload.get("copy_task_id"), creative_task_id=payload.get("creative_task_id"), reused=True)

    record = existing or BusinessRecord(record_type="proposal_revision", external_id=external_id, title=f"Обновлённое КП для {copy['client_name']}", status="draft", source="proposal_studio")
    if existing is None:
        db.add(record)
        db.flush()
    directory = Path(settings.document_storage_path) / "proposal-revisions"
    directory.mkdir(parents=True, exist_ok=True)
    stem = f"{external_id.lower()}-{_safe_filename(str(copy['client_name']))}"
    docx_path = directory / f"{stem}.docx"
    pdf_path = directory / f"{stem}.pdf"
    _build_docx(docx_path, copy)
    _build_pdf(pdf_path, copy)
    os.chmod(docx_path, 0o600)
    os.chmod(pdf_path, 0o600)
    files = {
        "docx": {"filename": docx_path.name, "storage_path": str(docx_path), "checksum_sha256": hashlib.sha256(docx_path.read_bytes()).hexdigest()},
        "pdf": {"filename": pdf_path.name, "storage_path": str(pdf_path), "checksum_sha256": hashlib.sha256(pdf_path.read_bytes()).hexdigest()},
    }
    from .platform import approval_engine

    approval = approval_engine.request(
        db,
        "proposal_release",
        "proposal_revision",
        str(record.id),
        "proposal_studio",
        {"external_id": external_id, "file_checksums": {kind: item["checksum_sha256"] for kind, item in files.items()}},
        "Проект КП должен быть проверен владельцем до отправки заказчику.",
    )
    record.status = "ready"
    record.data = {
        "source_filename": str(payload.get("source_filename") or source.name),
        "source_checksum_sha256": copy["source_checksum_sha256"],
        "copy_provider": copy["provider"],
        "creative_provider": design["provider"],
        "design_preset": design["preset"],
        "files": files,
        "approval_id": approval.id,
        "owner_approval_required_before_sending": True,
        "sent_to_client": False,
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }
    db.flush()
    return _revision_result(record, copy_task_id=payload.get("copy_task_id"), creative_task_id=payload.get("creative_task_id"), reused=False)


def _revision_result(record: BusinessRecord, *, copy_task_id: Any, creative_task_id: Any, reused: bool) -> dict[str, Any]:
    files = record.data.get("files") or {}
    return {
        "status": "ready_for_owner_review",
        "proposal_revision_id": record.id,
        "proposal_number": record.external_id,
        "approval_id": record.data.get("approval_id"),
        "owner_approval_required_before_sending": True,
        "sent_to_client": False,
        "reused": reused,
        "download_urls": {kind: f"/api/proposal-revisions/{record.id}/files/{kind}" for kind in files},
        "copy_task_id": copy_task_id,
        "creative_task_id": creative_task_id,
        "evidence": [
            {"type": "proposal_copy_review", "task_id": copy_task_id, "provider": record.data.get("copy_provider")},
            {"type": "proposal_design_review", "task_id": creative_task_id, "provider": record.data.get("creative_provider"), "preset": record.data.get("design_preset")},
            *[
                {"type": "document_export", "format": kind, "checksum_sha256": item.get("checksum_sha256")}
                for kind, item in files.items()
                if isinstance(item, dict)
            ],
        ],
    }
