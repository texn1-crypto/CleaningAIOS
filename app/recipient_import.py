from __future__ import annotations

import io
import re
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


EMAIL_PATTERN = re.compile(
    r"(?<![\w.+-])[A-Za-z0-9.!#$%&'*+/=?^_`{|}~-]+@[A-Za-z0-9-]+"
    r"(?:\.[A-Za-z0-9-]+)+(?![\w.-])"
)
SUPPORTED_RECIPIENT_SUFFIXES = {".xlsx", ".xlsm", ".docx", ".pdf"}
MAX_UNCOMPRESSED_OFFICE_BYTES = 50_000_000
MAX_EXTRACTED_CHARACTERS = 5_000_000
MAX_EXCEL_CELLS = 200_000
MAX_PDF_PAGES = 500


def _validate_office_archive(content: bytes) -> None:
    try:
        with ZipFile(io.BytesIO(content)) as archive:
            if sum(item.file_size for item in archive.infolist()) > MAX_UNCOMPRESSED_OFFICE_BYTES:
                raise ValueError("Office document expands beyond the safe processing limit")
    except BadZipFile as exc:
        raise ValueError("Invalid Office document") from exc


def _emails(text: str) -> set[str]:
    return {match.group(0).lower() for match in EMAIL_PATTERN.finditer(text or "")}


def _xlsx_emails(content: bytes) -> set[str]:
    _validate_office_archive(content)
    try:
        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    except Exception as exc:
        raise ValueError("Excel document could not be read") from exc
    result: set[str] = set()
    cells = 0
    characters = 0
    try:
        canonical_sheet = next(
            (
                sheet
                for sheet in workbook.worksheets
                if " ".join(sheet.title.lower().replace("ё", "е").split())
                in {"для импорта", "import", "import data", "данные для импорта"}
            ),
            None,
        )
        # A structured owner workbook may also contain raw/unlinked contact
        # sheets. Those are provenance, never an implicit recipient source.
        sheets = [canonical_sheet] if canonical_sheet is not None else workbook.worksheets
        for sheet in sheets:
            for row in sheet.iter_rows(values_only=True):
                for value in row:
                    cells += 1
                    if cells > MAX_EXCEL_CELLS:
                        raise ValueError("Excel document has too many cells")
                    if value is None:
                        continue
                    text = str(value)
                    characters += len(text)
                    if characters > MAX_EXTRACTED_CHARACTERS:
                        raise ValueError("Excel document contains too much text")
                    result.update(_emails(text))
    finally:
        workbook.close()
    return result


def _docx_emails(content: bytes) -> set[str]:
    _validate_office_archive(content)
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise ValueError("Word document could not be read") from exc
    result: set[str] = set()
    characters = 0
    for paragraph in document.paragraphs:
        characters += len(paragraph.text)
        if characters > MAX_EXTRACTED_CHARACTERS:
            raise ValueError("Word document contains too much text")
        result.update(_emails(paragraph.text))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                characters += len(cell.text)
                if characters > MAX_EXTRACTED_CHARACTERS:
                    raise ValueError("Word document contains too much text")
                result.update(_emails(cell.text))
    return result


def _pdf_emails(content: bytes) -> set[str]:
    try:
        reader = PdfReader(io.BytesIO(content))
        if reader.is_encrypted and reader.decrypt("") == 0:
            raise ValueError("Password-protected PDF is not supported")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise ValueError("PDF has too many pages")
        result: set[str] = set()
        characters = 0
        for page in reader.pages:
            text = page.extract_text() or ""
            characters += len(text)
            if characters > MAX_EXTRACTED_CHARACTERS:
                raise ValueError("PDF contains too much text")
            result.update(_emails(text))
        return result
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError("PDF document could not be read") from exc


def extract_recipient_emails(filename: str, content: bytes) -> list[str]:
    """Extract unique email addresses from a bounded, locally parsed document."""
    suffix = Path(filename or "").suffix.lower()
    if suffix not in SUPPORTED_RECIPIENT_SUFFIXES:
        raise ValueError("Only XLSX, XLSM, DOCX or searchable PDF files are supported")
    if not content:
        raise ValueError("Recipient document is empty")
    if suffix in {".xlsx", ".xlsm"}:
        result = _xlsx_emails(content)
    elif suffix == ".docx":
        result = _docx_emails(content)
    else:
        result = _pdf_emails(content)
    return sorted(result)
