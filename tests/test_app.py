import asyncio
from pathlib import Path

import pytest


def test_health(client):
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json()["database"] == "ok"


def test_task_and_agent_flow(client):
    created = client.post("/api/tasks", json={"title": "Проверить тендер", "agent_type": "tender"})
    assert created.status_code == 201
    task_id = created.json()["id"]
    result = client.post(f"/api/tasks/{task_id}/run")
    assert result.status_code == 200
    assert result.json()["status"] == "done"
    assert result.json()["result"]["submission_requires_owner_approval"] is True

    transitions = client.get(f"/api/tasks/{task_id}/transitions").json()
    assert [(row["from_status"], row["to_status"], row["reason"]) for row in transitions] == [
        ("", "open", "task_created"),
        ("open", "running", "agent_execution_started"),
        ("running", "done", "execution_evidence_verified"),
    ]
    assert client.post(f"/api/tasks/{task_id}/run").status_code == 409


def test_failed_task_retry_has_explicit_transition_history(client, monkeypatch):
    from app.agents import AGENTS

    class FailingAgent:
        name = "transition_test"

        def execute(self, db, payload):
            raise RuntimeError("controlled retry failure")

    monkeypatch.setitem(AGENTS, "transition_test", FailingAgent())
    task = client.post("/api/tasks", json={
        "title": "Transition retry test",
        "agent_type": "transition_test",
        "max_attempts": 2,
    }).json()
    first = client.post(f"/api/tasks/{task['id']}/run")
    assert first.status_code == 200
    assert first.json()["status"] == "queued"
    transitions = client.get(f"/api/tasks/{task['id']}/transitions").json()
    assert [(row["from_status"], row["to_status"]) for row in transitions] == [
        ("", "open"),
        ("open", "running"),
        ("running", "failed"),
        ("failed", "queued"),
    ]
    assert transitions[-1]["reason"] == "retry_scheduled"


def test_default_orchestrator_task_runs(client):
    task = client.post("/api/tasks", json={"title": "Общая операционная задача", "payload": {"message": "Проверить объект"}}).json()
    result = client.post(f"/api/tasks/{task['id']}/run").json()
    assert result["status"] == "done"
    assert result["result"]["coordinated"] is True


def test_owner_approval_gate(client):
    task = client.post("/api/tasks", json={"title": "Подать заявку", "agent_type": "tender", "payload": {"action_kind": "tender_submission"}}).json()
    blocked = client.post(f"/api/tasks/{task['id']}/run").json()
    assert blocked["status"] == "blocked"
    assert blocked["result"]["reason"] == "owner_approval_required"
    approval_id = blocked["result"]["approval_id"]
    approved = client.post(f"/api/approvals/{approval_id}/approve", json={"note": "Проверено владельцем"})
    assert approved.status_code == 200
    approval_history = client.get(f"/api/tasks/{task['id']}/transitions").json()
    assert approval_history[-1]["to_status"] == "queued"
    assert approval_history[-1]["reason"] == "owner_approval_granted"
    completed = client.post(f"/api/tasks/{task['id']}/run").json()
    assert completed["status"] == "done"


def test_company_brain_versions_knowledge(client):
    payload = {"namespace": "company", "key": "service_area", "value": {"city": "Москва"}, "confidence": 0.9}
    assert client.put("/api/brain", json=payload).json()["version"] == 1
    payload["value"] = {"city": "Москва", "radius_km": 50}
    assert client.put("/api/brain", json=payload).json()["version"] == 2
    snapshot = client.get("/api/brain?namespace=company").json()
    assert snapshot["company.service_area"]["value"]["radius_km"] == 50


def test_record_emits_domain_event(client):
    from uuid import UUID

    record = client.post("/api/records", json={"record_type": "lead", "title": "УК Север"})
    assert record.status_code == 201
    events = client.get("/api/events", headers={"X-Role": "manager"}).json()
    event = next(x for x in events if x["event_type"] == "lead.created" and x["aggregate_id"] == str(record.json()["id"]))
    assert str(UUID(event["event_id"])) == event["event_id"]
    assert event["schema_version"] == 1
    assert event["correlation_id"] == event["event_id"]
    assert event["causation_id"] == ""
    assert event["actor"] == "api-user"
    assert event["occurred_at"]
    assert event["deliveries"] == []


def test_sales_lifecycle_contacts_and_summary(client):
    lead = client.post("/api/records", json={"record_type": "lead", "title": "БЦ Восток", "status": "new", "data": {"budget": 120000}}).json()
    changed = client.patch(f"/api/records/{lead['id']}", json={"status": "qualified", "data": {"next_action": "send_proposal"}})
    assert changed.status_code == 200
    assert changed.json()["data"] == {"budget": 120000, "next_action": "send_proposal"}
    touch = client.post(f"/api/records/{lead['id']}/contacts", json={"channel": "phone", "direction": "outbound", "outcome": "meeting_booked"})
    assert touch.status_code == 201
    assert len(client.get(f"/api/records/{lead['id']}/contacts").json()) == 1
    sales = client.get("/api/modules/summary").json()["sales"]
    assert sales["qualified"] >= 1
    assert sales["pipeline_amount"] >= 120000


def test_domain_rules_prevent_incomplete_records(client):
    invalid_type = client.post("/api/records", json={"record_type": "Invalid Type", "title": "Невалидный тип"})
    assert invalid_type.status_code == 422
    lost = client.post("/api/records", json={"record_type": "lead", "title": "Неполный лид", "status": "lost"})
    assert lost.status_code == 422
    finance = client.post("/api/records", json={"record_type": "expense", "title": "Без суммы", "status": "pending"})
    assert finance.status_code == 422
    tender = client.post("/api/records", json={"record_type": "tender", "title": "Без срока", "status": "preparing"})
    assert tender.status_code == 422


def test_event_bus_routes_domain_work_to_agent(client):
    from app.db import SessionLocal
    from app.platform import process_next_event

    lead = client.post("/api/records", json={"record_type": "lead", "title": "Маршрутизируемый лид"}).json()
    with SessionLocal() as db:
        while process_next_event(db):
            pass
    tasks = client.get("/api/tasks").json()
    assert any(x["agent_type"] == "sales" and x["payload"].get("record_id") == str(lead["id"]) for x in tasks)


def test_event_bus_consumer_receipt_prevents_duplicate_routing(client):
    from sqlalchemy import func, select

    from app.db import SessionLocal
    from app.models import DomainEvent, EventConsumerReceipt, Task
    from app.platform import event_bus, process_next_event

    aggregate_id = "receipt-idempotency-test"
    with SessionLocal() as db:
        event = event_bus.publish(
            db,
            "lead.created",
            "lead",
            aggregate_id,
            {"title": "Receipt test"},
            idempotency_key="test:event-receipt-idempotency",
            actor="test-suite",
            correlation_id="correlation-receipt-test",
        )
        db.commit()
        event_db_id = event.id
        event_uid = event.event_id

        while True:
            processed = process_next_event(db)
            current = db.get(DomainEvent, event_db_id)
            if current and current.status == "published":
                break
            assert processed is not None

        receipt = db.scalar(select(EventConsumerReceipt).where(
            EventConsumerReceipt.event_id == event_db_id,
            EventConsumerReceipt.consumer == "domain_router",
        ))
        assert receipt is not None
        assert receipt.status == "succeeded"
        assert receipt.attempts == 1
        assert receipt.result_ref.startswith("task:")
        routed_before = db.scalar(select(func.count(Task.id)).where(Task.payload["event_uid"].as_string() == event_uid))

        current.status = "pending"
        db.commit()
        process_next_event(db)
        routed_after = db.scalar(select(func.count(Task.id)).where(Task.payload["event_uid"].as_string() == event_uid))
        db.refresh(receipt)
        assert routed_after == routed_before == 1
        assert receipt.attempts == 1

    api_event = next(row for row in client.get("/api/events", headers={"X-Role": "manager"}).json() if row["event_id"] == event_uid)
    assert len(api_event["deliveries"]) == 1
    delivery = api_event["deliveries"][0]
    assert delivery["consumer"] == "domain_router"
    assert delivery["status"] == "succeeded"
    assert delivery["attempts"] == 1
    assert delivery["result_ref"].startswith("task:")
    assert delivery["last_error"] == ""
    assert delivery["processed_at"]


def test_event_bus_persists_failed_consumer_delivery(monkeypatch):
    from pydantic import ValidationError
    from sqlalchemy import select

    from app import platform
    from app.db import SessionLocal
    from app.models import DomainEvent, EventConsumerReceipt

    with SessionLocal() as db:
        while platform.process_next_event(db):
            pass
        with pytest.raises(ValidationError):
            platform.event_bus.publish(db, "Invalid Event", "lead", "invalid")
        db.rollback()

        event = platform.event_bus.publish(
            db,
            "lead.created",
            "lead",
            "receipt-failure-test",
            idempotency_key="test:event-receipt-failure",
        )
        db.commit()
        event_db_id = event.id

        def fail_route(db, event):
            raise RuntimeError("consumer route failed")

        monkeypatch.setattr(platform, "route_event", fail_route)
        with pytest.raises(RuntimeError, match="consumer route failed"):
            platform.process_next_event(db)

        failed_event = db.get(DomainEvent, event_db_id)
        receipt = db.scalar(select(EventConsumerReceipt).where(
            EventConsumerReceipt.event_id == event_db_id,
            EventConsumerReceipt.consumer == "domain_router",
        ))
        assert failed_event.status == "pending"
        assert failed_event.attempts == 1
        assert receipt is not None
        assert receipt.status == "failed"
        assert receipt.attempts == 1
        assert receipt.last_error == "consumer route failed"


def test_company_graph_economics_and_simulator(client):
    customer = client.post("/api/entities", json={"entity_type": "client", "name": "УК Тест"}).json()
    site = client.post("/api/entities", json={"entity_type": "site", "name": "ЖК Тест", "parent_id": customer["id"], "data": {"materials_cost": 10000, "logistics_cost": 5000}}).json()
    contract = client.post("/api/entities", json={"entity_type": "contract", "name": "Контракт ЖК", "parent_id": site["id"], "data": {"monthly_revenue": 200000}})
    assert contract.status_code == 201
    client.post("/api/entities", json={"entity_type": "shift", "name": "Дневная смена", "parent_id": site["id"], "data": {"payroll_cost": 90000, "employee_id": 1}})
    graph = client.get("/api/company/graph").json()
    assert any(x["id"] == customer["id"] and x["children"][0]["id"] == site["id"] for x in graph["roots"])
    economics = client.get(f"/api/finance/site-economics?site_id={site['id']}").json()[0]
    assert economics["profit"] == 95000
    simulation = client.post("/api/simulations", json={"site_id": site["id"], "payroll_change_percent": 10}).json()
    assert simulation["base"]["profit"] == 86000


def test_goals_structured_decisions_and_outcomes(client):
    goal = client.post("/api/goals", headers={"X-Role": "manager"}, json={"title": "Рост прибыли", "metric": "monthly_profit", "baseline": 100, "target": 200}).json()
    progress = client.patch(f"/api/goals/{goal['id']}/progress", json={"current": 150}).json()
    assert progress["progress_percent"] == 50
    decision = client.post("/api/structured-decisions", json={"title": "Увеличить резерв", "problem": "Незакрытые смены", "confidence": 0.87, "options": [{"id": "A"}]}).json()
    outcome = client.put(f"/api/decisions/{decision['id']}/outcome", headers={"X-Role": "manager"}, json={"successful": True, "actual_value": 42}).json()
    assert outcome["successful"] is True


def test_tender_scoring_and_document_registry(client):
    tender = client.post("/api/records", json={"record_type": "tender", "title": "Клининг БЦ", "data": {"expected_margin": 80, "company_fit": 90, "competition_risk": 20, "contract_risk": 20, "logistics_fit": 90, "staffing_fit": 80, "strategic_value": 60}}).json()
    score = client.post(f"/api/tenders/{tender['id']}/score").json()
    assert score["score"] >= 80
    document = client.post(f"/api/tenders/{tender['id']}/documents", json={"name": "ТЗ.pdf", "source_url": "https://example.invalid/spec.pdf", "analysis": {"penalties": "1% per day"}})
    assert document.status_code == 201
    assert document.json()["status"] == "analyzed"


def test_csv_import_and_bulk_campaign_approval(client, monkeypatch):
    import base64
    from app.config import settings
    for field in ("smtp_host", "smtp_username", "smtp_password", "smtp_from_email"):
        monkeypatch.setattr(settings, field, "")
    content = base64.b64encode("company,email,budget\nУК Альфа,alpha@example.com,500000\n,missing@example.com,10\n".encode()).decode()
    imported = client.post("/api/imports/leads", json={"filename": "leads.csv", "content_base64": content}).json()
    assert imported["imported_rows"] == 1
    for address in ("alpha@example.com", "second@example.com"):
        consent = client.put("/api/outreach/consents", json={
            "address": address,
            "source_url": "https://consent.example.test/evidence",
            "evidence": f"Documented opt-in from {address}",
        })
        assert consent.status_code == 200
    launch = {"campaign_key": "approved-campaign", "recipients": ["alpha@example.com", "second@example.com"], "subject": "Клининг", "body": "Предложение"}
    blocked = client.post("/api/outreach/campaigns/launch", headers={"X-Role": "manager"}, json=launch).json()
    assert blocked["status"] == "waiting_approval"
    client.post(f"/api/approvals/{blocked['approval_id']}/approve", json={"note": "Разрешаю"})
    launch["approval_id"] = blocked["approval_id"]
    queued = client.post("/api/outreach/campaigns/launch", headers={"X-Role": "manager"}, json=launch).json()
    assert queued["status"] == "credentials_required"
    assert queued["queued"] == 2


def test_bulk_campaign_approval_cannot_authorize_changed_content(client):
    assert client.put("/api/outreach/consents", json={
        "address": "one@example.com",
        "source_url": "https://consent.example.test/one",
        "evidence": "Documented opt-in for tamper test",
    }).status_code == 200
    launch = {"campaign_key": "tamper-proof-campaign", "recipients": ["one@example.com"], "subject": "Original", "body": "Original body"}
    blocked = client.post("/api/outreach/campaigns/launch", headers={"X-Role": "manager"}, json=launch).json()
    client.post(f"/api/approvals/{blocked['approval_id']}/approve", json={"note": "Original approved"})
    changed = {**launch, "subject": "Changed", "approval_id": blocked["approval_id"]}
    second = client.post("/api/outreach/campaigns/launch", headers={"X-Role": "manager"}, json=changed).json()
    assert second["status"] == "waiting_approval"
    assert second["approval_id"] != blocked["approval_id"]


def test_xlsx_lead_import(client):
    import base64
    import io
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["company", "email", "district"])
    sheet.append(["ТСЖ Excel", "excel@example.com", "ЮЗАО"])
    stream = io.BytesIO(); workbook.save(stream)
    payload = {"filename": "leads.xlsx", "content_base64": base64.b64encode(stream.getvalue()).decode()}
    result = client.post("/api/imports/leads", json=payload)
    assert result.status_code == 201
    assert result.json()["imported_rows"] == 1


def test_import_size_limit_and_mailbox_secret_reference(client, monkeypatch):
    import base64
    from app.config import settings

    monkeypatch.setattr(settings, "max_import_bytes", 3)
    oversized = client.post("/api/imports/leads", json={"filename": "large.csv", "content_base64": base64.b64encode(b"four").decode()})
    assert oversized.status_code == 413
    invalid_secret = client.post("/api/outreach/mailboxes", json={"name": "Unsafe mailbox", "address": "safe@example.com", "secret_ref": "API_KEY"})
    assert invalid_secret.status_code == 422


def test_structured_decision_creates_bound_approval(client):
    decision = client.post("/api/structured-decisions", json={"title": "Подписать договор", "problem": "Новый объект", "requires_approval": True, "approval_kind": "contract"}).json()
    assert decision["approval_id"] is not None
    approved = client.post(f"/api/approvals/{decision['approval_id']}/approve", json={"note": "Проверено"})
    assert approved.status_code == 200
    decisions = client.get("/api/decisions").json()
    assert next(x for x in decisions if x["id"] == decision["id"])["status"] == "approved"


def test_integration_status_is_truthful(client):
    status = client.get("/api/integrations").json()
    assert status["tender_sources"]["status"] in {"configured", "source_configuration_required"}
    assert status["llm"]["status"] in {
        "configured",
        "credentials_required",
        "model_configuration_required",
        "provider_configuration_required",
        "version_configuration_required",
    }
    assert status["llm"]["provider"] == "multi_provider_advisory_router"
    assert set(status["llm"]["providers"]) == {"openai_responses", "anthropic_messages"}


def test_llm_adapter_uses_structured_responses_contract(monkeypatch):
    import json
    from app.config import settings
    from app import llm

    captured = {}

    class Response:
        def raise_for_status(self): return None
        def json(self):
            output = {
                "summary": "Стабильное состояние",
                "risks": ["Один риск"],
                "data_gaps": ["Нет данных источника"],
                "recommendations": [{"title": "Проверить маржу", "agent_type": "finance", "rationale": "Маржа требует проверки", "priority": "high", "needs_owner_decision": False}],
            }
            return {"status": "completed", "model": "test-model", "output": [{"type": "message", "content": [{"type": "output_text", "text": json.dumps(output)}]}], "usage": {"total_tokens": 42}}

    class Client:
        def __init__(self, *args, **kwargs): captured["headers"] = kwargs["headers"]
        def __enter__(self): return self
        def __exit__(self, *args): return False
        def post(self, url, json): captured.update({"url": url, "payload": json}); return Response()

    monkeypatch.setattr(llm.httpx, "Client", Client)
    monkeypatch.setattr(settings, "llm_api_key", "test-secret")
    monkeypatch.setattr(settings, "llm_base_url", "https://api.example/v1")
    monkeypatch.setattr(settings, "llm_model", "test-model")
    result = llm.llm_advisor.review({"business_health": 90})
    assert result["status"] == "succeeded"
    assert result["recommendations"][0]["agent_type"] == "finance"
    assert captured["url"] == "https://api.example/v1/responses"
    assert captured["payload"]["text"]["format"]["strict"] is True
    assert captured["payload"]["store"] is False
    assert captured["headers"]["Authorization"] == "Bearer test-secret"


def test_anthropic_adapter_uses_native_messages_structured_output(monkeypatch):
    import json
    from app.config import settings
    from app import llm

    captured = {}

    class Response:
        def raise_for_status(self): return None
        def json(self):
            output = {
                "summary": "Claude: устойчивое состояние",
                "risks": [],
                "data_gaps": ["Нет данных о марже"],
                "recommendations": [{
                    "title": "Проверить экономику объекта",
                    "agent_type": "finance",
                    "rationale": "Маржа не рассчитана",
                    "priority": "high",
                    "needs_owner_decision": False,
                }],
            }
            return {
                "model": "claude-test",
                "stop_reason": "end_turn",
                "content": [{"type": "text", "text": json.dumps(output)}],
                "usage": {"input_tokens": 40, "output_tokens": 20},
            }

    class Client:
        def __init__(self, *args, **kwargs): captured["headers"] = kwargs["headers"]
        def __enter__(self): return self
        def __exit__(self, *args): return False
        def post(self, url, json): captured.update({"url": url, "payload": json}); return Response()

    monkeypatch.setattr(llm.httpx, "Client", Client)
    monkeypatch.setattr(settings, "anthropic_api_key", "anthropic-test-secret")
    monkeypatch.setattr(settings, "anthropic_base_url", "https://api.anthropic.example")
    monkeypatch.setattr(settings, "anthropic_model", "claude-test")
    result = llm.AnthropicMessagesAdvisor().review({"business_health": 88})
    assert result["status"] == "succeeded"
    assert result["provider"] == "anthropic_messages"
    assert result["recommendations"][0]["agent_type"] == "finance"
    assert captured["url"] == "https://api.anthropic.example/v1/messages"
    assert captured["payload"]["output_config"]["format"]["type"] == "json_schema"
    portable_schema = captured["payload"]["output_config"]["format"]["schema"]
    assert "maxItems" not in json.dumps(portable_schema)
    assert portable_schema["additionalProperties"] is False
    assert captured["payload"]["system"] == llm.SYSTEM_PROMPT
    assert captured["headers"]["x-api-key"] == "anthropic-test-secret"
    assert captured["headers"]["anthropic-version"] == "2023-06-01"
    assert "Authorization" not in captured["headers"]


def test_multi_provider_router_assigns_tasks_and_falls_back(monkeypatch):
    from app.config import settings
    from app import llm

    monkeypatch.setattr(settings, "llm_provider", "auto")
    monkeypatch.setattr(settings, "llm_api_key", "openai-secret")
    monkeypatch.setattr(settings, "anthropic_api_key", "anthropic-secret")
    router = llm.LLMAdvisor()
    monkeypatch.setattr(router.anthropic, "review", lambda snapshot: {
        "status": "succeeded", "provider": "anthropic_messages", "model": "claude-test", "recommendations": [],
    })
    monkeypatch.setattr(router.openai, "analyze_request", lambda *args: {
        "status": "succeeded", "provider": "openai_responses", "model": "gpt-test", "should_create_improvement": False,
    })
    business = router.review({"business_health": 90})
    request = router.analyze_request("проверь", {"kind": "task"}, {"classification": "capability_gap"})
    assert business["provider"] == "anthropic_messages"
    assert business["attempted_providers"] == ["anthropic_messages"]
    assert request["provider"] == "openai_responses"
    assert request["attempted_providers"] == ["openai_responses"]

    monkeypatch.setattr(router.anthropic, "review", lambda snapshot: {
        "status": "unavailable", "provider": "anthropic_messages", "model": "claude-test", "recommendations": [],
    })
    monkeypatch.setattr(router.openai, "review", lambda snapshot: {
        "status": "succeeded", "provider": "openai_responses", "model": "gpt-test", "recommendations": [],
    })
    fallback = router.review({"business_health": 90})
    assert fallback["provider"] == "openai_responses"
    assert fallback["attempted_providers"] == ["anthropic_messages", "openai_responses"]


def test_ai_ceo_falls_back_without_llm_credentials(client, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "llm_api_key", "")
    task = client.post("/api/tasks", json={"title": "CEO fallback check", "agent_type": "ceo"}).json()
    result = client.post(f"/api/tasks/{task['id']}/run").json()["result"]
    assert result["llm_advice"]["status"] == "credentials_required"
    assert isinstance(result["business_health"], int)


def test_ai_ceo_creates_only_safe_llm_tasks(client, monkeypatch):
    from app.agents import llm_advisor

    monkeypatch.setattr(llm_advisor, "review", lambda snapshot: {
        "status": "succeeded",
        "model": "test-model",
        "summary": "Test",
        "risks": [],
        "data_gaps": [],
        "recommendations": [
            {"title": "LLM safe finance analysis unique", "agent_type": "finance", "rationale": "Проверить данные", "priority": "high", "needs_owner_decision": False},
            {"title": "LLM protected commitment unique", "agent_type": "finance", "rationale": "Оплатить счет", "priority": "high", "needs_owner_decision": True},
        ],
    })
    task = client.post("/api/tasks", json={"title": "CEO LLM task check", "agent_type": "ceo"}).json()
    result = client.post(f"/api/tasks/{task['id']}/run").json()["result"]
    created = {x["title"]: x for x in result["tasks_created"]}
    assert created["LLM safe finance analysis unique"]["source"] == "llm_advisory"
    assert "LLM protected commitment unique" not in created
    queued = {x["title"]: x for x in client.get("/api/tasks").json()}
    assert queued["LLM safe finance analysis unique"]["payload"]["advisory_only"] is True


def test_research_agent_runs_real_collector_contract(client, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "tender_sources", "")
    task = client.post("/api/tasks", json={"title": "Collect configured tender feeds", "agent_type": "research"}).json()
    result = client.post(f"/api/tasks/{task['id']}/run").json()
    assert result["status"] == "done"
    assert result["result"]["collection"] == "tenders"
    assert result["result"]["status"] == "source_configuration_required"
    assert result["result"]["created"] == 0


def test_mission_control_renders(client, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "company_name", "CleaningAIOS")
    response = client.get("/")
    assert response.status_code == 200
    assert "Порядок, который работает на ваш бизнес" in response.text
    assert 'content="http://testserver/static/og-cleaningaios.png"' in response.text
    assert "CleaningAIOS — профессиональный клининг" in response.text
    assert "__COMPANY_NAME__" not in response.text
    assert client.get("/mission-control").status_code == 200
    assert "Mission Control" in client.get("/mission-control").text


def test_public_site_has_real_multipage_catalog_prices_and_sitemap(client):
    from app.site_pages import PRICE_ROWS, SERVICE_DETAILS, SERVICE_IMAGE_SIZES, _price

    for path in ("/services", "/prices", "/about", "/contacts", "/journal"):
        response = client.get(path)
        assert response.status_code == 200
        assert "CleaningAIOS" in response.text
        assert "h2oclean" not in response.text.lower()
    catalog = client.get("/services").text
    assert catalog.count('class="catalog-card reveal"') == len(SERVICE_DETAILS)
    assert catalog.count('loading="lazy"') == len(SERVICE_DETAILS)
    assert "/static/service-imagery.css" in catalog
    assert len({row["image"] for row in SERVICE_DETAILS.values()}) >= 7
    for image, dimensions in SERVICE_IMAGE_SIZES.items():
        assert dimensions[0] > 0 and dimensions[1] > 0
        assert Path("app" + image).is_file()
    for slug in SERVICE_DETAILS:
        detail = client.get(f"/services/{slug}")
        assert detail.status_code == 200
        assert 'fetchpriority="high"' in detail.text
    assert client.get("/services/unknown-service").status_code == 404

    homepage = client.get("/").text
    assert "/static/services/business-center-lobby-v1.jpg" in homepage
    assert 'fetchpriority="high"' in homepage

    price_page = client.get("/prices").text
    for name, general, regular, after in PRICE_ROWS:
        assert name in price_page
        for reference in (general, regular, after):
            price = _price(reference)
            assert reference * 0.95 <= price < reference
            assert f"от {price} ₽" in price_page
    sitemap = client.get("/sitemap.xml")
    assert sitemap.status_code == 200
    assert "/services/business-centers" in sitemap.text
    assert "/prices" in sitemap.text


def test_contacts_page_exposes_configured_click_to_call_phone(client, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "company_phone", "+7 995 599-60-95")
    response = client.get("/contacts")
    assert response.status_code == 200
    assert "+7 995 599-60-95" in response.text
    assert 'href="tel:+79955996095"' in response.text


def test_public_site_exposes_only_safe_social_profile_urls(client, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "social_telegram_url", "https://t.me/cleaningaios")
    monkeypatch.setattr(settings, "social_vk_url", "javascript:alert(1)")
    monkeypatch.setattr(settings, "social_odnoklassniki_url", "https://user:secret@ok.ru/group")
    monkeypatch.setattr(settings, "social_instagram_url", "http://instagram.com/cleaningaios")
    payload = client.get("/api/public/site").json()
    assert payload["company"]["social"] == {
        "telegram": "https://t.me/cleaningaios",
        "vk": "",
        "odnoklassniki": "",
        "instagram": "",
    }


def test_unified_inbox_content_and_operations_views(client):
    lead = client.post("/api/records", json={"record_type": "lead", "title": "Лид из inbox"}).json()
    message = client.post("/api/inbox", json={"channel": "email", "external_id": "mail-001", "sender": "client@example.com", "subject": "Запрос цены", "record_id": lead["id"]})
    assert message.status_code == 201
    assert client.post("/api/inbox", json={"channel": "email", "external_id": "mail-001"}).status_code == 409
    updated = client.patch(f"/api/inbox/{message.json()['id']}", json={"status": "assigned", "record_id": lead["id"]}).json()
    assert updated["status"] == "assigned"
    content = client.post("/api/marketing/content", json={"channel": "telegram", "title": "Преимущества уборки", "status": "draft"})
    assert content.status_code == 201
    assert len(client.get("/api/marketing/content?status=draft").json()) == 1


def test_staffing_vacancy_draft_payment_calendar_and_quality(client):
    customer = client.post("/api/entities", json={"entity_type": "client", "name": "Клиент процессов"}).json()
    site = client.post("/api/entities", json={"entity_type": "site", "name": "Объект процессов", "parent_id": customer["id"]}).json()
    vacancy = client.post("/api/entities", json={"entity_type": "vacancy", "name": "Уборщик в БЦ", "data": {"district": "САО", "schedule": "2/2", "rate": 3500}}).json()
    client.post("/api/entities", json={"entity_type": "shift", "name": "Незакрытая смена", "parent_id": site["id"], "data": {"payroll_cost": 3500}})
    client.post("/api/entities", json={"entity_type": "complaint", "name": "Жалоба на качество", "parent_id": site["id"], "status": "open", "data": {"sla_deadline": "2020-01-01T00:00:00+00:00"}})
    staffing = client.get("/api/hr/staffing").json()
    assert len(staffing["unfilled_shifts"]) >= 1
    assert "САО" in client.get(f"/api/hr/vacancies/{vacancy['id']}/telegram-draft").json()["text"]
    client.post("/api/records", json={"record_type": "payment", "title": "Платёж клиента", "status": "overdue", "data": {"amount": 10000}, "deadline_at": "2020-01-01T00:00:00"})
    assert client.get("/api/finance/payment-calendar").json()[0]["overdue"] is True
    assert client.get("/api/operations/quality").json()["sla_breached"] >= 1


def test_tender_feed_and_document_download(client, monkeypatch, tmp_path):
    from app.config import settings
    from app import integrations

    feed = {"items": [{"external_id": "feed-1", "title": "Уборка бизнес-центра", "deadline_at": "2030-01-01T12:00:00Z", "data": {"expected_margin": 80, "company_fit": 90}, "documents": [{"name": "contract.pdf", "url": "https://feed.example/contract.pdf", "content_type": "application/pdf"}]}]}

    class Response:
        headers = {"content-type": "application/pdf", "content-length": "7"}
        def raise_for_status(self): return None
        def json(self): return feed
        def iter_bytes(self): yield b"PDFDATA"
        def __enter__(self): return self
        def __exit__(self, *args): return False

    class Client:
        def __init__(self, *args, **kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *args): return False
        def get(self, url): return Response()
        def stream(self, method, url): return Response()

    monkeypatch.setattr(integrations.httpx, "Client", Client)
    monkeypatch.setattr(integrations.socket, "getaddrinfo", lambda *args, **kwargs: [(2, 1, 6, "", ("93.184.216.34", 443))])
    monkeypatch.setattr(settings, "tender_sources", "https://feed.example/tenders")
    monkeypatch.setattr(settings, "document_storage_path", str(tmp_path))
    collected = client.post("/api/tender-sources/collect", headers={"X-Role": "manager"}).json()
    assert collected["created"] == 1
    from app.db import SessionLocal
    from app.models import TenderDocument
    from sqlalchemy import select
    with SessionLocal() as db:
        document_id = db.scalar(select(TenderDocument.id).where(TenderDocument.source_url == "https://feed.example/contract.pdf"))
    downloaded = client.post(f"/api/tender-documents/{document_id}/download").json()
    assert downloaded["status"] == "downloaded"
    assert downloaded["bytes"] == 7


def test_delivery_event_suppresses_bounced_recipient(client):
    queued = client.post("/api/outreach/messages", json={"campaign_key": "bounce-test", "recipient": "bounce@example.com", "subject": "Test", "body": "Body"}).json()
    event = client.post("/api/outreach/delivery-events", json={"event_type": "bounce", "recipient": "bounce@example.com", "message_id": queued["id"], "reason": "mailbox unavailable"})
    assert event.status_code == 201
    assert event.json()["suppressed"] is True
    retry = client.post("/api/outreach/messages", json={"campaign_key": "bounce-retry", "recipient": "bounce@example.com", "subject": "Test", "body": "Body"})
    assert retry.status_code == 409


def test_worker_sends_real_attachment(client, monkeypatch):
    import base64
    from sqlalchemy import update
    from app.config import settings
    from app.db import SessionLocal
    from app.models import OutboundMessage
    from app import worker

    sent = []
    class SMTP:
        def __init__(self, *args, **kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *args): return False
        def starttls(self): pass
        def login(self, username, password): pass
        def send_message(self, message): sent.append(message)

    monkeypatch.setattr(worker.smtplib, "SMTP", SMTP)
    for field, value in {"smtp_host": "smtp.example", "smtp_port": 587, "smtp_username": "user", "smtp_password": "secret", "smtp_from_email": "sender@example.com"}.items(): monkeypatch.setattr(settings, field, value)
    with SessionLocal() as db:
        db.execute(update(OutboundMessage).where(OutboundMessage.status.in_(["queued", "waiting_configuration"])).values(status="sent"))
        row = OutboundMessage(campaign_key="attachment-test", recipient="attach@example.com", subject="Документ", body="Смотрите вложение", attachments=[{"filename": "offer.txt", "content_type": "text/plain", "content_base64": base64.b64encode(b"offer").decode()}])
        db.add(row); db.commit(); row_id = row.id
        assert worker.send_next_email(db) is True
        db.refresh(row); assert row.status == "sent"
    assert sent and sent[0].is_multipart()
    assert any(part.get_filename() == "offer.txt" for part in sent[0].walk())
    assert "token=" in sent[0].get_body(preferencelist=("plain",)).get_content()


def test_worker_supports_smtp_implicit_tls_on_port_465(client, monkeypatch):
    from sqlalchemy import update

    from app import worker
    from app.config import settings
    from app.db import SessionLocal
    from app.models import OutboundMessage

    sent = []

    class SMTPSSL:
        def __init__(self, *args, **kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *args): return False
        def login(self, username, password): return None
        def send_message(self, message): sent.append(message)

    monkeypatch.setattr(worker.smtplib, "SMTP_SSL", SMTPSSL)
    monkeypatch.setattr(settings, "smtp_host", "smtp.example")
    monkeypatch.setattr(settings, "smtp_port", 465)
    monkeypatch.setattr(settings, "smtp_username", "implicit-tls@example.com")
    monkeypatch.setattr(settings, "smtp_password", "test-secret")
    monkeypatch.setattr(settings, "smtp_from_email", "implicit-tls@example.com")
    with SessionLocal() as db:
        db.execute(update(OutboundMessage).where(OutboundMessage.status.in_(["queued", "waiting_configuration"])).values(status="sent"))
        row = OutboundMessage(
            campaign_key="implicit-tls-test",
            recipient="implicit-tls-recipient@example.com",
            subject="TLS",
            body="TLS body",
        )
        db.add(row)
        db.commit()
        assert worker.send_next_email(db) is True
        db.refresh(row)
        assert row.status == "sent"
    assert len(sent) == 1


def test_outreach_delivery_window_opens_at_nine_moscow(monkeypatch):
    from datetime import datetime

    from app import worker
    from app.config import settings

    monkeypatch.setattr(settings, "outreach_timezone", "Europe/Moscow")
    monkeypatch.setattr(settings, "outreach_daily_start_hour", 9)

    before_start, next_start = worker.outreach_delivery_window(datetime(2026, 8, 13, 5, 59))
    assert before_start is None
    assert next_start == datetime(2026, 8, 13, 6, 0)

    window_start, next_start = worker.outreach_delivery_window(datetime(2026, 8, 13, 6, 0))
    assert window_start == datetime(2026, 8, 13, 6, 0)
    assert next_start == datetime(2026, 8, 14, 6, 0)

    overnight_start, next_start = worker.outreach_delivery_window(datetime(2026, 8, 13, 22, 0))
    assert overnight_start is None
    assert next_start == datetime(2026, 8, 14, 6, 0)


def test_worker_reports_each_daily_delivery_and_stops_at_limit(client, monkeypatch):
    from datetime import datetime

    from sqlalchemy import select, update

    from app import notifications, worker
    from app.config import settings
    from app.db import SessionLocal
    from app.models import OutboundMessage, OwnerNotification

    sent = []
    telegram = []

    class SMTP:
        def __init__(self, *args, **kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *args): return False
        def starttls(self): return None
        def login(self, username, password): return None
        def send_message(self, message): sent.append(message["To"])

    class TelegramResponse:
        def raise_for_status(self): return None

    class TelegramClient:
        def __init__(self, *args, **kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *args): return False
        def post(self, url, json):
            telegram.append(json["text"])
            return TelegramResponse()

    monkeypatch.setattr(worker.smtplib, "SMTP", SMTP)
    monkeypatch.setattr(notifications.httpx, "Client", TelegramClient)
    for field, value in {
        "smtp_host": "smtp.example",
        "smtp_port": 587,
        "smtp_username": "sender@example.com",
        "smtp_password": "secret",
        "smtp_from_email": "sender@example.com",
        "owner_telegram_id": "999",
        "telegram_bot_token": "123456:test-token",
        "outreach_timezone": "Europe/Moscow",
        "outreach_daily_start_hour": 9,
        "outreach_per_minute": 10,
        "outreach_per_day": 2,
    }.items():
        monkeypatch.setattr(settings, field, value)

    now = datetime(2026, 8, 13, 6, 0)
    campaign_key = "daily-progress-test"
    with SessionLocal() as db:
        db.execute(update(OutboundMessage).values(status="sent", sent_at=None))
        db.add_all([
            OutboundMessage(
                campaign_key=campaign_key,
                recipient=f"daily-{index}@example.com",
                subject="Daily progress",
                body="Body",
                status="queued",
                scheduled_at=now,
            )
            for index in range(1, 4)
        ])
        db.commit()

        assert worker.send_next_email(db, now=now) is True
        assert worker.send_next_email(db, now=now) is True
        assert worker.send_next_email(db, now=now) is False

        progress = db.scalars(
            select(OwnerNotification)
            .where(OwnerNotification.resource_id == campaign_key)
            .order_by(OwnerNotification.id)
        ).all()
        assert [row.subject for row in progress] == [
            "📨 Рассылка: 1/2",
            "📨 Рассылка: 2/2",
        ]
        assert "Всего по файлу: 1/3" in progress[0].body
        assert "Всего по файлу: 2/3" in progress[1].body
        assert "14.08.2026 09:00" in progress[1].body
        assert len({row.idempotency_key for row in progress}) == 2
        db.execute(
            update(OwnerNotification)
            .where(OwnerNotification.resource_id != campaign_key)
            .values(status="sent")
        )
        db.commit()
        assert notifications.send_next_owner_notification(db) is True
        assert notifications.send_next_owner_notification(db) is True
        assert telegram[0].startswith("📨 Рассылка: 1/2")
        assert telegram[1].startswith("📨 Рассылка: 2/2")
        remaining = db.scalar(
            select(OutboundMessage.status).where(
                OutboundMessage.campaign_key == campaign_key,
                OutboundMessage.recipient == "daily-3@example.com",
            )
        )
        assert remaining == "queued"
        db.execute(
            update(OutboundMessage)
            .where(OutboundMessage.campaign_key == campaign_key)
            .values(status="sent", sent_at=None)
        )
        db.commit()

    assert sent == ["daily-1@example.com", "daily-2@example.com"]


def test_outreach_suppression_and_deduplication(client):
    assert client.post("/api/outreach/suppress", json={"address": "stop@example.com"}).status_code == 201
    blocked = client.post("/api/outreach/messages", json={"campaign_key": "c1", "recipient": "stop@example.com", "subject": "Hi", "body": "Body"})
    assert blocked.status_code == 409
    payload = {"campaign_key": "c1", "recipient": "ok@example.com", "subject": "Hi", "body": "Body"}
    assert client.post("/api/outreach/messages", json=payload).status_code == 201
    assert client.post("/api/outreach/messages", json=payload).status_code == 409


def test_rbac_rejects_viewer_write(client):
    response = client.post("/api/tasks", headers={"X-Role": "viewer"}, json={"title": "Denied"})
    assert response.status_code == 403


def test_production_rbac_derives_role_from_api_key(monkeypatch):
    from fastapi import HTTPException
    from app.config import settings
    from app.security import principal, validate_production_security

    monkeypatch.setattr(settings, "environment", "production")
    monkeypatch.setattr(settings, "api_key", "owner-secret")
    monkeypatch.setattr(settings, "manager_api_key", "manager-secret")
    monkeypatch.setattr(settings, "operator_api_key", "operator-secret")
    monkeypatch.setattr(settings, "viewer_api_key", "viewer-secret")
    monkeypatch.setattr(settings, "telegram_callback_secret", "callback-secret")
    validate_production_security()
    authenticated = principal(x_api_key="manager-secret", x_actor="spoofed-owner", x_role="owner")
    assert authenticated.role == "manager"
    assert authenticated.subject == "manager-api-key"
    with pytest.raises(HTTPException) as exc:
        principal(x_api_key="unknown", x_actor="api-user", x_role="owner")
    assert exc.value.status_code == 401


def test_tender_url_guard_blocks_private_networks(monkeypatch):
    from fastapi import HTTPException
    from app import integrations

    with pytest.raises(HTTPException):
        integrations._safe_url("http://127.0.0.1/admin")
    monkeypatch.setattr(integrations.socket, "getaddrinfo", lambda *args, **kwargs: [(2, 1, 6, "", ("10.0.0.5", 443))])
    with pytest.raises(HTTPException):
        integrations._safe_url("https://feed.example/tenders")


def test_unsubscribe_link_requires_valid_signature(client):
    from app.security import unsubscribe_token

    address = "signed-unsubscribe@example.com"
    assert client.get(f"/api/outreach/unsubscribe?email={address}&token=invalid").status_code == 403
    response = client.get(f"/api/outreach/unsubscribe?email={address}&token={unsubscribe_token(address)}")
    assert response.status_code == 200
    assert response.json()["unsubscribed"] is True


def test_telegram_requires_owner_id(monkeypatch):
    from app.config import settings
    from app.bot import build_application

    monkeypatch.setattr(settings, "telegram_bot_token", "123456:fake-token-for-startup-check")
    monkeypatch.setattr(settings, "owner_telegram_id", "")
    with pytest.raises(RuntimeError, match="OWNER_TELEGRAM_ID is empty"):
        build_application()


def test_russian_chat_reads_existing_sections():
    from app.chat import understand_russian_message

    assert understand_russian_message("Меню") == {"kind": "menu"}
    assert understand_russian_message("Открой меню") == {"kind": "menu"}
    assert understand_russian_message("Покажи текущие задачи")["kind"] == "tasks"
    assert understand_russian_message("Что с тендерами?")["record_type"] == "tender"
    assert understand_russian_message("Покажи финансы")["module"] == "finance"
    assert understand_russian_message("Как дела у системы?")["kind"] == "dashboard"
    assert understand_russian_message("Пришли отчет о проделанной работе")["kind"] == "activity_report"
    assert understand_russian_message("Запусти весь функционал чат бота")["kind"] == "system_self_check"
    assert understand_russian_message("Сколько нужно времени на выполнение задачи?") == {
        "kind": "task_eta",
        "task_id": None,
    }
    assert understand_russian_message("Когда будет готова задача 42?") == {
        "kind": "task_eta",
        "task_id": 42,
    }


def test_menu_request_is_supported_and_does_not_create_an_improvement(client, monkeypatch):
    from app.chat import understand_russian_message
    from app.config import settings

    monkeypatch.setattr(settings, "llm_api_key", "")
    message = "Меню"
    response = client.post(
        "/api/request-analysis",
        json={"message": message, "intent": understand_russian_message(message)},
    )
    assert response.status_code == 200
    analysis = response.json()
    assert analysis["classification"] == "supported"
    assert analysis["fully_supported"] is True
    assert analysis["improvement_id"] is None


def test_telegram_menu_phrase_opens_interactive_menu(monkeypatch):
    from app import bot

    calls = []

    async def fake_api(method, path, **kwargs):
        calls.append((method, path, kwargs))
        assert path == "/api/request-analysis"
        return {"classification": "supported", "improvement_id": None}

    class Message:
        text = "Меню"
        reply_to_message = None

        def __init__(self):
            self.replies = []

        async def reply_text(self, value, **kwargs):
            self.replies.append((value, kwargs))

    class User:
        id = 123

    class Update:
        effective_message = Message()
        effective_user = User()

    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    asyncio.run(bot.natural_language(Update(), None))

    assert len(calls) == 1
    text, kwargs = Update.effective_message.replies[-1]
    assert text == "CleaningAI OS · выберите раздел:"
    buttons = [button for row in kwargs["reply_markup"].inline_keyboard for button in row]
    assert any(button.callback_data == "outreach" for button in buttons)
    assert any(button.callback_data == "tasks" for button in buttons)


def test_russian_chat_routes_misspelled_social_setup_request_to_marketing(client, monkeypatch):
    from app.chat import understand_russian_message
    from app.config import settings

    monkeypatch.setattr(settings, "llm_api_key", "")
    message = "начните офрмлять социалбные сети вк и однокласники"
    intent = understand_russian_message(message)
    assert intent["kind"] == "task"
    assert intent["agent_type"] == "marketing"
    assert intent["payload"]["action"] == "prepare_social_account_setup"
    assert intent["payload"]["channels"] == ["vk", "odnoklassniki"]
    analysis = client.post("/api/request-analysis", json={"message": message, "intent": intent}).json()
    assert analysis["classification"] == "supported"
    assert analysis["improvement_id"] is None


def test_social_account_setup_persists_real_progress_and_audits_failures(client, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "social_vk_url", "")
    monkeypatch.setattr(settings, "vk_community_id", "")
    monkeypatch.setattr(settings, "vk_community_token", "")
    monkeypatch.setattr(settings, "social_odnoklassniki_url", "")
    monkeypatch.setattr(settings, "odnoklassniki_group_id", "")
    monkeypatch.setattr(settings, "odnoklassniki_application_key", "")
    monkeypatch.setattr(settings, "odnoklassniki_session_secret", "")
    task = client.post("/api/tasks", json={
        "title": "Начать оформление VK и Одноклассников",
        "agent_type": "marketing",
        "payload": {"action": "prepare_social_account_setup", "channels": ["vk", "odnoklassniki"]},
        "max_attempts": 1,
    }).json()
    completed = client.post(f"/api/tasks/{task['id']}/run").json()
    assert completed["status"] == "done"
    result = completed["result"]
    assert result["status"] == "setup_in_progress"
    assert result["records_created"] == 2
    assert result["external_accounts_created"] == 0
    assert result["publication_started"] is False
    assert {row["channel"] for row in result["platforms"]} == {"vk", "odnoklassniki"}
    assert all(row["status"] == "integration_configuration_required" for row in result["platforms"])
    records = client.get("/api/records?record_type=social_account_setup").json()
    assert {row["data"]["channel"] for row in records} >= {"vk", "odnoklassniki"}
    assert any(
        row["action"] == "task.completed" and row["resource_id"] == str(task["id"])
        for row in client.get("/api/audit").json()
    )

    failed_task = client.post("/api/tasks", json={
        "title": "Неподдерживаемая социальная площадка",
        "agent_type": "marketing",
        "payload": {"action": "prepare_social_account_setup", "channels": ["unknown"]},
        "max_attempts": 1,
    }).json()
    failed = client.post(f"/api/tasks/{failed_task['id']}/run").json()
    assert failed["status"] == "failed"
    assert "Unsupported social channels" in failed["result"]["error"]
    assert any(
        row["action"] == "task.failed" and row["resource_id"] == str(failed_task["id"])
        for row in client.get("/api/audit").json()
    )


def test_telegram_social_setup_runs_agent_and_returns_actual_result(monkeypatch):
    from app import bot

    calls = []

    async def fake_api(method, path, **kwargs):
        calls.append((method, path, kwargs))
        if path == "/api/request-analysis":
            return {"classification": "supported", "improvement_id": None}
        if path == "/api/tasks":
            return {"id": 611, "agent_type": "marketing", "title": "Соцсети"}
        if path == "/api/tasks/611/run":
            return {
                "status": "done",
                "result": {
                    "status": "setup_in_progress",
                    "records_created": 2,
                    "records_updated": 0,
                    "external_accounts_created": 0,
                    "platforms": [
                        {"channel": "vk", "status": "integration_configuration_required", "missing_configuration": ["SOCIAL_VK_URL"]},
                        {"channel": "odnoklassniki", "status": "integration_configuration_required", "missing_configuration": ["SOCIAL_ODNOKLASSNIKI_URL"]},
                    ],
                },
            }
        raise AssertionError(path)

    class Message:
        text = "начните офрмлять социалбные сети вк и однокласники"

        def __init__(self):
            self.replies = []

        async def reply_text(self, value, **kwargs):
            self.replies.append(value)

    class User:
        id = 123

    class Update:
        effective_message = Message()
        effective_user = User()

    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    asyncio.run(bot.natural_language(Update(), None))
    reply = Update.effective_message.replies[-1]
    assert "Оформление социальных сетей начато" in reply
    assert "VK" in reply and "Одноклассники" in reply
    assert "Внешних аккаунтов автоматически создано: 0" in reply
    task_payload = next(kwargs["json"] for _, path, kwargs in calls if path == "/api/tasks")
    assert task_payload["agent_type"] == "marketing"
    assert task_payload["payload"]["action"] == "prepare_social_account_setup"
    assert task_payload["max_attempts"] == 1


def test_telegram_social_setup_preserves_failed_status(monkeypatch):
    from app import bot

    async def fake_api(method, path, **kwargs):
        if path == "/api/request-analysis":
            return {"classification": "supported", "improvement_id": None}
        if path == "/api/tasks":
            return {"id": 612, "agent_type": "marketing", "title": "Соцсети"}
        if path == "/api/tasks/612/run":
            return {"status": "failed", "result": {"error": "adapter failure"}}
        raise AssertionError(path)

    class Message:
        text = "начните оформлять социальные сети вк и одноклассники"

        def __init__(self):
            self.replies = []

        async def reply_text(self, value, **kwargs):
            self.replies.append(value)

    class User:
        id = 123

    class Update:
        effective_message = Message()
        effective_user = User()

    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    asyncio.run(bot.natural_language(Update(), None))
    reply = Update.effective_message.replies[-1]
    assert "не началось" in reply
    assert "adapter failure" in reply


def test_task_timing_report_does_not_treat_empty_agent_summary_as_completion(client):
    from app.chat import understand_russian_message

    message = (
        "Собери базу управляющих компаний из всех возможных источников, "
        "найди сайты и почты, проверь их и выгрузи PDF, XLSX и Word"
    )
    intent = understand_russian_message("Сколько нужно времени на выполнение задачи?")
    assessment = client.post(
        "/api/request-analysis",
        json={"message": "Сколько нужно времени на выполнение задачи?", "intent": intent},
    ).json()
    assert assessment["classification"] == "supported"
    assert assessment["improvement_id"] is None

    business_task = client.post("/api/tasks", json={
        "title": message,
        "agent_type": "marketing",
        "payload": {
            "source": "telegram_natural_language",
            "original_message": message,
        },
    }).json()
    quality_gated = client.post(f"/api/tasks/{business_task['id']}/run").json()
    assert quality_gated["status"] == "blocked"
    assert quality_gated["result"]["improvement_id"]
    assert quality_gated["result"]["ceo_incident_task_id"]

    timing_task = client.post("/api/tasks", json={
        "title": "Оценить срок выполнения задачи",
        "agent_type": "orchestrator",
        "payload": {
            "action": "task_timing_report",
            "task_id": None,
            "source": "telegram_read_request",
        },
        "max_attempts": 1,
    }).json()
    completed = client.post(f"/api/tasks/{timing_task['id']}/run").json()
    result = completed["result"]

    assert completed["status"] == "done"
    assert result["task"]["id"] == business_task["id"]
    assert result["timing_status"] == "blocked"
    assert result["result_verified"] is False
    assert result["remaining_seconds"] is None
    assert result["actual_runtime_seconds"] is not None
    assert result["planning_estimate"]["min_hours"] == 8
    assert result["planning_estimate"]["max_hours"] == 24
    assert any(
        row["action"] == "task.completed" and row["resource_id"] == str(timing_task["id"])
        for row in client.get("/api/audit").json()
    )


def test_telegram_task_timing_returns_truthful_result(monkeypatch):
    from app import bot

    calls = []

    async def fake_api(method, path, **kwargs):
        calls.append((method, path, kwargs))
        if path == "/api/request-analysis":
            return {"classification": "supported", "improvement_id": None}
        if path == "/api/tasks":
            return {"id": 601}
        if path == "/api/tasks/601/run":
            return {
                "status": "done",
                "result": {
                    "outcome": "completed",
                    "task": {"id": 9, "title": "Собрать базу УК", "status": "done"},
                    "timing_status": "result_unverified",
                    "actual_runtime_seconds": 0.013,
                    "planning_estimate": {
                        "min_hours": 8,
                        "max_hours": 24,
                        "confidence": "low",
                        "starts_after": ["подключены источники"],
                    },
                    "reason": "Подтвержденный результат отсутствует.",
                },
            }
        raise AssertionError(path)

    class Message:
        text = "Сколько нужно времени на выполнение задачи?"

        def __init__(self):
            self.replies = []

        async def reply_text(self, value, **kwargs):
            self.replies.append(value)

    class User:
        id = 123

    class Update:
        effective_message = Message()
        effective_user = User()

    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    asyncio.run(bot.natural_language(Update(), None))

    reply = Update.effective_message.replies[-1]
    assert "Задача #9" in reply
    assert "результата или файла нет" in reply
    assert "8–24 ч" in reply
    assert "это не равно" in reply
    task_payload = next(kwargs["json"] for _, path, kwargs in calls if path == "/api/tasks")
    assert task_payload["payload"] == {
        "action": "task_timing_report",
        "task_id": None,
        "source": "telegram_read_request",
    }
    assert task_payload["max_attempts"] == 1


def test_system_self_check_is_safe_audited_and_truthful(client, monkeypatch):
    import json

    from app.chat import understand_russian_message
    from app.config import settings

    monkeypatch.setattr(settings, "telegram_bot_token", "test-secret-must-not-leak")
    monkeypatch.setattr(settings, "owner_telegram_id", "123")
    monkeypatch.setattr(settings, "smtp_password", "smtp-secret-must-not-leak")
    monkeypatch.setattr(settings, "llm_api_key", "")
    message = "Запусти весь функционал чат бота"
    intent = understand_russian_message(message)
    analysis = client.post("/api/request-analysis", json={"message": message, "intent": intent}).json()
    assert analysis["classification"] == "supported"
    assert analysis["improvement_id"] is None
    approvals_before = len(client.get("/api/approvals").json())

    task = client.post("/api/tasks", json={
        "title": "Безопасная самопроверка функционала чат-бота",
        "agent_type": "orchestrator",
        "payload": {"action": "system_self_check"},
        "max_attempts": 1,
    }).json()
    completed = client.post(f"/api/tasks/{task['id']}/run").json()
    result = completed["result"]

    assert completed["status"] == "done"
    assert result["outcome"] == "completed"
    assert result["check_kind"] == "system_functional_readiness"
    assert result["overall_status"] in {"ready", "partial"}
    assert result["summary"]["total"] == len(result["checks"])
    assert result["safety"] == {
        "protected_actions_executed": False,
        "external_messages_sent": False,
        "financial_commitments_created": False,
        "owner_approval_bypassed": False,
    }
    assert {row["name"] for row in result["checks"]} >= {
        "PostgreSQL", "API и Orchestrator", "Telegram-бот", "Sales/CRM",
        "Тендеры и Research", "Email и горячие лиды", "Публичный сайт и лид-форма",
    }
    serialized = json.dumps(result, ensure_ascii=False)
    assert "test-secret-must-not-leak" not in serialized
    assert "smtp-secret-must-not-leak" not in serialized
    assert len(client.get("/api/approvals").json()) == approvals_before
    assert any(item["type"] == "database_probe" for item in result["evidence"])
    assert any(
        row["action"] == "task.completed" and row["resource_id"] == str(task["id"])
        for row in client.get("/api/audit").json()
    )


def test_telegram_system_self_check_returns_result_not_task_acceptance(monkeypatch):
    from app import bot

    calls = []

    async def fake_api(method, path, **kwargs):
        calls.append((method, path, kwargs))
        if path == "/api/request-analysis":
            return {"classification": "supported", "improvement_id": None}
        if path == "/api/tasks":
            return {"id": 501, "agent_type": "orchestrator", "title": "Самопроверка"}
        if path == "/api/tasks/501/run":
            return {
                "status": "done",
                "result": {
                    "outcome": "completed",
                    "overall_status": "partial",
                    "summary": {"ready": 8, "total": 14},
                    "checks": [
                        {"name": "PostgreSQL", "status": "ready", "detail": "Контрольный запрос выполнен."},
                        {"name": "Email", "status": "credentials_required", "detail": "SMTP не настроен."},
                    ],
                    "credentials_required": ["SMTP_PASSWORD"],
                },
            }
        raise AssertionError(path)

    class Message:
        text = "Запусти весь функционал чат бота"

        def __init__(self):
            self.replies = []

        async def reply_text(self, value, **kwargs):
            self.replies.append(value)

    class User:
        id = 123

    class Update:
        effective_message = Message()
        effective_user = User()

    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    asyncio.run(bot.natural_language(Update(), None))
    reply = Update.effective_message.replies[-1]
    assert "частично готово" in reply
    assert "PostgreSQL" in reply
    assert "Платежи, договоры" in reply
    assert "Task accepted" not in reply
    task_payload = next(kwargs["json"] for method, path, kwargs in calls if path == "/api/tasks")
    assert task_payload["payload"]["action"] == "system_self_check"
    assert task_payload["max_attempts"] == 1


def test_telegram_system_self_check_preserves_failed_status(monkeypatch):
    from app import bot

    async def fake_api(method, path, **kwargs):
        if path == "/api/request-analysis":
            return {"classification": "supported", "improvement_id": None}
        if path == "/api/tasks":
            return {"id": 502}
        if path == "/api/tasks/502/run":
            return {"status": "failed", "result": {"error": "database unavailable"}}
        raise AssertionError(path)

    class Message:
        text = "Запусти весь функционал чат бота"

        def __init__(self):
            self.replies = []

        async def reply_text(self, value, **kwargs):
            self.replies.append(value)

    class User:
        id = 123

    class Update:
        effective_message = Message()
        effective_user = User()

    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    asyncio.run(bot.natural_language(Update(), None))
    assert "не завершилась" in Update.effective_message.replies[-1]
    assert "failed" in Update.effective_message.replies[-1]


def test_activity_report_is_a_real_audited_orchestrator_result(client, monkeypatch):
    from app.chat import understand_russian_message
    from app.config import settings

    monkeypatch.setattr(settings, "llm_api_key", "")
    message = "Пришли отчёт о проделанной работе"
    intent = understand_russian_message(message)
    analysis = client.post("/api/request-analysis", json={"message": message, "intent": intent}).json()
    assert analysis["classification"] == "supported"
    assert analysis["improvement_id"] is None
    active_statuses = {"open", "queued", "running"}
    active_before = sum(row["status"] in active_statuses for row in client.get("/api/tasks").json())

    task = client.post("/api/tasks", json={
        "title": "Сформировать отчёт о проделанной работе",
        "agent_type": "orchestrator",
        "payload": {"action": "system_activity_report", "period_hours": 24},
        "max_attempts": 1,
    }).json()
    completed = client.post(f"/api/tasks/{task['id']}/run").json()
    result = completed["result"]
    assert completed["status"] == "done"
    assert result["outcome"] == "completed"
    assert result["report_kind"] == "system_activity"
    assert result["summary"]["tasks_active"] == active_before
    assert isinstance(result["recent_completed_tasks"], list)
    assert any(item["type"] == "database_snapshot" for item in result["evidence"])
    audit_rows = client.get("/api/audit").json()
    assert any(
        row["action"] == "task.completed" and row["resource_id"] == str(task["id"])
        for row in audit_rows
    )


def test_telegram_activity_report_returns_actual_result(monkeypatch):
    from app import bot

    calls = []

    async def fake_api(method, path, **kwargs):
        calls.append((method, path, kwargs))
        if path == "/api/request-analysis":
            return {"classification": "supported", "improvement_id": None}
        if path == "/api/tasks":
            return {"id": 401, "agent_type": "orchestrator", "title": "Отчёт"}
        if path == "/api/tasks/401/run":
            return {
                "status": "done",
                "result": {
                    "outcome": "completed",
                    "period_hours": 24,
                    "summary": {
                        "tasks_completed": 3,
                        "tasks_active": 0,
                        "tasks_failed": 0,
                        "tasks_blocked": 0,
                        "queued_improvements": 1,
                        "pending_approvals": 0,
                    },
                    "recent_completed_tasks": [
                        {"id": 9, "agent_type": "sales", "title": "Подготовить КП"}
                    ],
                    "blockers": [],
                },
            }
        raise AssertionError(path)

    class Message:
        text = "Пришли отчет о проделанной работе"
        replies = []

        async def reply_text(self, value, **kwargs):
            self.replies.append(value)

    class User:
        id = 123

    class Update:
        effective_message = Message()
        effective_user = User()

    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    asyncio.run(bot.natural_language(Update(), None))
    assert "Выполнено задач: 3" in Update.effective_message.replies[-1]
    assert "#9 [sales] Подготовить КП" in Update.effective_message.replies[-1]
    task_payload = next(kwargs["json"] for method, path, kwargs in calls if path == "/api/tasks")
    assert task_payload["payload"]["action"] == "system_activity_report"
    assert task_payload["max_attempts"] == 1


def test_scheduled_activity_report_is_30_minutes_and_notification_is_idempotent(client):
    from uuid import uuid4

    from app.db import SessionLocal
    from app.models import OwnerNotification
    from app.reports import format_activity_report
    from sqlalchemy import func, select

    notification_key = f"scheduled-report-test:{uuid4()}"
    payload = {
        "action": "system_activity_report",
        "period_minutes": 30,
        "source": "scheduler",
        "notify_owner": True,
        "scheduled_window_start": "2026-08-13T00:00:00",
        "notification_idempotency_key": notification_key,
    }
    first = client.post(
        "/api/tasks",
        json={"title": "Регулярный отчёт тест", "agent_type": "orchestrator", "payload": payload},
    ).json()
    first_result = client.post(f"/api/tasks/{first['id']}/run").json()["result"]
    second = client.post(
        "/api/tasks",
        json={"title": "Повтор регулярного отчёта", "agent_type": "orchestrator", "payload": payload},
    ).json()
    second_result = client.post(f"/api/tasks/{second['id']}/run").json()["result"]

    assert first_result["period_minutes"] == 30
    assert first_result["period_hours"] == 0.5
    assert first_result["owner_notification"] in {"queued", "waiting_configuration"}
    assert second_result["owner_notification_id"] == first_result["owner_notification_id"]
    assert "за 30 мин." in format_activity_report(first_result)
    with SessionLocal() as db:
        assert db.scalar(
            select(func.count()).select_from(OwnerNotification).where(
                OwnerNotification.idempotency_key == notification_key
            )
        ) == 1


def test_scheduler_creates_one_owner_report_per_window(monkeypatch):
    from sqlalchemy import create_engine, select
    from sqlalchemy.orm import sessionmaker
    from sqlalchemy.pool import StaticPool

    from app import scheduler
    from app.db import Base
    from app.models import Task

    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    session_factory = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    monkeypatch.setattr(scheduler, "SessionLocal", session_factory)
    monkeypatch.setattr(scheduler.settings, "owner_activity_report_interval_minutes", 30)

    scheduler.schedule_cycle()
    scheduler.schedule_cycle()

    with session_factory() as db:
        all_tasks = db.scalars(select(Task)).all()
        reports = db.scalars(
            select(Task).where(Task.title.like("Регулярный отчёт владельцу · %"))
        ).all()
        development = [
            row for row in all_tasks if row.payload.get("origin") == "ceo_continuous_backlog"
        ]
        assert len(reports) == 1
        assert len(development) == 4
        assert {row.payload["scope"] for row in development} == {
            "website",
            "sales",
            "marketing",
            "system",
        }
        assert reports[0].payload["period_minutes"] == 30
        assert reports[0].payload["notify_owner"] is True
        assert reports[0].payload["notification_idempotency_key"].startswith(
            "owner-activity-report:"
        )


def test_ceo_keeps_safe_deduplicated_development_backlog():
    from datetime import datetime, timedelta

    from sqlalchemy import create_engine, select
    from sqlalchemy.orm import sessionmaker
    from sqlalchemy.pool import StaticPool

    from app.db import Base
    from app.models import Task
    from app.operations import maintain_ceo_development_backlog
    from app.task_state import transition_task

    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    session_factory = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    now = datetime(2026, 8, 13, 0, 0)

    with session_factory() as db:
        first = maintain_ceo_development_backlog(db, now=now, cadence_hours=24)
        duplicate = maintain_ceo_development_backlog(db, now=now, cadence_hours=24)
        db.commit()

        assert len(first) == 4
        assert duplicate == []
        assert {row.payload["scope"] for row in first} == {
            "website",
            "sales",
            "marketing",
            "system",
        }
        assert all(row.payload["advisory_only"] is True for row in first)
        assert all(row.payload["external_actions_require_owner_approval"] is True for row in first)

        for row in first:
            transition_task(db, row, "running", actor=row.agent_type, reason="test_execution")
            transition_task(db, row, "done", actor="orchestrator", reason="test_completed")
        second = maintain_ceo_development_backlog(db, now=now, cadence_hours=24)
        db.commit()

        assert len(second) == 4
        assert all(row.status == "queued" for row in second)
        assert all(row.run_after == now + timedelta(hours=24) for row in second)
        assert len(db.scalars(select(Task)).all()) == 8


def test_russian_chat_routes_business_requests_to_agents():
    from app.chat import understand_russian_message

    research = understand_russian_message("Найди тендеры по уборке бизнес-центров")
    assert research["kind"] == "task"
    assert research["agent_type"] == "research"
    assert research["payload"]["collection"] == "tenders"

    sales = understand_russian_message("Создай задачу связаться с новым клиентом")
    assert sales["agent_type"] == "sales"
    assert sales["payload"]["source"] == "telegram_natural_language"


def test_russian_chat_resolves_improve_this_from_replied_text():
    from app.chat import understand_russian_message

    missing = understand_russian_message("улучши это")
    assert missing["kind"] == "clarification"
    assert "Ответьте" in missing["message"]

    intent = understand_russian_message(
        "улучши это",
        referenced_text="мы хотим предложить вам качественно и быстро убрать ваш офис",
    )
    assert intent["kind"] == "task"
    assert intent["agent_type"] == "copywriter"
    assert intent["payload"]["action"] == "improve_referenced_text"
    assert intent["payload"]["external_send"] is False


def test_copywriter_improves_replied_text_with_audited_draft(client):
    from app.chat import understand_russian_message

    intent = understand_russian_message(
        "улучши это",
        referenced_text="мы хотим предложить вам качественно и быстро убрать ваш офис",
    )
    analysis = client.post(
        "/api/request-analysis",
        json={"message": "улучши это", "intent": intent, "source_channel": "telegram"},
    ).json()
    assert analysis["classification"] == "supported"
    assert analysis["improvement_id"] is None

    task = client.post("/api/tasks", json={
        "title": intent["title"],
        "agent_type": intent["agent_type"],
        "payload": intent["payload"],
        "max_attempts": 1,
    }).json()
    completed = client.post(f"/api/tasks/{task['id']}/run").json()
    assert completed["status"] == "done"
    assert completed["result"]["status"] == "ready"
    assert "Предлагаем" in completed["result"]["improved_text"]
    assert "подготовим расчёт" in completed["result"]["improved_text"]
    assert completed["result"]["external_send"] is False
    assert completed["result"]["owner_review_required"] is True
    assert completed["result"]["evidence"][0]["type"] == "text_revision"
    audit = client.get("/api/audit").json()
    assert any(row["action"] == "task.completed" and row["resource_id"] == str(task["id"]) for row in audit)

    failed = client.post("/api/tasks", json={
        "title": "Missing reply context",
        "agent_type": "copywriter",
        "payload": {"action": "improve_referenced_text", "source": "api_test"},
        "max_attempts": 1,
    }).json()
    failed_result = client.post(f"/api/tasks/{failed['id']}/run").json()
    assert failed_result["status"] == "failed"
    assert "Referenced text is required" in failed_result["result"]["error"]


def test_telegram_improve_this_returns_real_draft(monkeypatch):
    from app import bot

    calls = []

    async def fake_api(method, path, **kwargs):
        calls.append((method, path, kwargs))
        if path == "/api/request-analysis":
            return {"classification": "supported", "improvement_id": None}
        if path == "/api/tasks":
            return {"id": 731, "agent_type": "copywriter", "title": "Улучшить текст"}
        if path == "/api/tasks/731/run":
            return {
                "status": "done",
                "result": {
                    "status": "ready",
                    "improved_text": "Предлагаем уборку по согласованному регламенту.",
                    "changes": ["Убраны разговорные формулировки."],
                    "external_send": False,
                },
            }
        raise AssertionError(path)

    class RepliedMessage:
        text = "мы хотим предложить вам качественно и быстро убрать ваш офис"
        caption = None

    class Message:
        text = "улучши это"
        reply_to_message = RepliedMessage()

        def __init__(self):
            self.replies = []

        async def reply_text(self, value, **kwargs):
            self.replies.append(value)

    class User:
        id = 123

    class Update:
        effective_message = Message()
        effective_user = User()

    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    asyncio.run(bot.natural_language(Update(), None))

    assert "Обновлённый черновик" in Update.effective_message.replies[-1]
    assert "никуда не отправлен" in Update.effective_message.replies[-1]
    task_payload = next(kwargs["json"] for _, path, kwargs in calls if path == "/api/tasks")
    assert task_payload["payload"]["action"] == "improve_referenced_text"
    assert task_payload["max_attempts"] == 1


def test_russian_chat_routes_feedback_on_previous_letter():
    from app.chat import understand_russian_message

    intent = understand_russian_message("дай обратную связь по моему предыдущему письму")
    assert intent["kind"] == "task"
    assert intent["agent_type"] == "copywriter"
    assert intent["payload"]["action"] == "review_previous_text"
    assert intent["payload"]["referenced_text"] == ""
    assert intent["payload"]["external_send"] is False


def test_previous_request_is_reviewed_with_audited_result(client):
    from app.chat import understand_russian_message

    original = "Добрый день! Мы хотим предложить вам уборку качественно и быстро. пароль: temporary-secret"
    first_intent = understand_russian_message("улучши это", referenced_text=original)
    first = client.post("/api/request-analysis", json={
        "message": original,
        "intent": first_intent,
        "source_channel": "telegram",
        "source_user": "feedback-owner",
    })
    assert first.status_code == 200

    feedback_intent = understand_russian_message("дай обратную связь по моему предыдущему письму")
    analysis = client.post("/api/request-analysis", json={
        "message": "дай обратную связь по моему предыдущему письму",
        "intent": feedback_intent,
        "source_channel": "telegram",
        "source_user": "feedback-owner",
    }).json()
    assert analysis["classification"] == "supported"
    assert analysis["context_found"] is True
    resolved = analysis["resolved_intent"]
    assert resolved["payload"]["action"] == "review_previous_text"
    assert "temporary-secret" not in resolved["payload"]["referenced_text"]
    assert "[REDACTED]" in resolved["payload"]["referenced_text"]

    task = client.post("/api/tasks", json={
        "title": resolved["title"],
        "agent_type": resolved["agent_type"],
        "payload": resolved["payload"],
        "max_attempts": 1,
    }).json()
    completed = client.post(f"/api/tasks/{task['id']}/run").json()
    result = completed["result"]
    assert completed["status"] == "done"
    assert result["status"] == "ready"
    assert "следующий шаг" in result["feedback_text"]
    assert "Предлагаем" in result["revised_text"]
    assert result["external_send"] is False
    assert result["owner_review_required"] is True
    assert result["evidence"][0]["type"] == "text_review"
    audit = client.get("/api/audit").json()
    assert any(row["action"] == "task.completed" and row["resource_id"] == str(task["id"]) for row in audit)

    missing = client.post("/api/tasks", json={
        "title": "Feedback without history",
        "agent_type": "copywriter",
        "payload": {"action": "review_previous_text"},
        "max_attempts": 1,
    }).json()
    failed = client.post(f"/api/tasks/{missing['id']}/run").json()
    assert failed["status"] == "failed"
    assert "Referenced text is required" in failed["result"]["error"]


def test_telegram_feedback_uses_resolved_previous_context(monkeypatch):
    from app import bot

    calls = []
    resolved_intent = {
        "kind": "task",
        "title": "Подготовить обратную связь по предыдущему тексту",
        "agent_type": "copywriter",
        "priority": "normal",
        "payload": {
            "action": "review_previous_text",
            "referenced_text": "Мы хотим предложить вам уборку качественно и быстро",
            "external_send": False,
        },
        "protected": False,
    }

    async def fake_api(method, path, **kwargs):
        calls.append((method, path, kwargs))
        if path == "/api/request-analysis":
            return {"classification": "supported", "resolved_intent": resolved_intent, "context_found": True}
        if path == "/api/tasks":
            return {"id": 841, "agent_type": "copywriter", "title": resolved_intent["title"]}
        if path == "/api/tasks/841/run":
            return {
                "status": "done",
                "result": {
                    "status": "ready",
                    "feedback_text": "• Не указан следующий шаг.",
                    "revised_text": "Предлагаем уборку по согласованному регламенту.",
                    "external_send": False,
                },
            }
        raise AssertionError(path)

    class Message:
        text = "дай обратную связь по моему предыдущему письму"
        reply_to_message = None

        def __init__(self):
            self.replies = []

        async def reply_text(self, value, **kwargs):
            self.replies.append(value)

    class User:
        id = 123

    class Update:
        effective_message = Message()
        effective_user = User()

    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    asyncio.run(bot.natural_language(Update(), None))

    reply = Update.effective_message.replies[-1]
    assert "Обратная связь" in reply
    assert "Предлагаемый вариант" in reply
    assert "никуда не отправлен" in reply
    task_payload = next(kwargs["json"] for _, path, kwargs in calls if path == "/api/tasks")
    assert task_payload["payload"]["action"] == "review_previous_text"
    assert task_payload["max_attempts"] == 1


def test_russian_chat_recognizes_proposal_client_without_command():
    from app.chat import understand_russian_message

    intent = understand_russian_message("Подготовь коммерческое предложение в PDF для тестового клиента Request Analyst")
    assert intent["kind"] == "task"
    assert intent["agent_type"] == "sales"
    assert intent["payload"]["action"] == "generate_proposal"
    assert intent["payload"]["client_query"] == "Request Analyst"


@pytest.mark.parametrize(
    ("message", "agent_type", "action_kind"),
    [
        ("Оплати счет поставщику", "finance", "financial"),
        ("Подпиши договор с новым клиентом", "sales", "contract"),
        ("Подай заявку на этот тендер", "tender", "tender_submission"),
        ("Найми уборщика на объект", "hr", "hr_final"),
        ("Разошли предложение всем клиентам", "sales", "bulk_outreach"),
    ],
)
def test_russian_chat_preserves_owner_approval_gates(message, agent_type, action_kind):
    from app.chat import understand_russian_message

    intent = understand_russian_message(message)
    assert intent["agent_type"] == agent_type
    assert intent["payload"]["action_kind"] == action_kind
    assert intent["protected"] is True


def test_russian_chat_protected_task_is_blocked_by_runtime(client):
    from app.chat import understand_russian_message

    intent = understand_russian_message("Оплати счет поставщику")
    task = client.post("/api/tasks", json={
        "title": intent["title"],
        "agent_type": intent["agent_type"],
        "priority": intent["priority"],
        "payload": intent["payload"],
    }).json()
    result = client.post(f"/api/tasks/{task['id']}/run").json()
    assert result["status"] == "blocked"
    assert result["result"]["reason"] == "owner_approval_required"


def test_telegram_application_registers_natural_language_handler(monkeypatch):
    from telegram.ext import CommandHandler, MessageHandler
    from app.config import settings
    from app.bot import build_application

    monkeypatch.setattr(settings, "telegram_bot_token", "123456:fake-token-for-startup-check")
    monkeypatch.setattr(settings, "owner_telegram_id", "123")
    application = build_application()
    handlers = [handler for group in application.handlers.values() for handler in group]
    assert any(isinstance(handler, MessageHandler) for handler in handlers)
    commands = {command for handler in handlers if isinstance(handler, CommandHandler) for command in handler.commands}
    assert {"outreach", "mailing", "cancel"}.issubset(commands)


def test_outreach_summary_is_owner_safe_and_manager_guarded(client):
    address = "bot-outreach-summary@example.com"
    assert client.put("/api/outreach/consents", json={
        "address": address,
        "source_url": "https://consent.example.test/bot-panel",
        "evidence": "Documented opt-in for the Telegram outreach panel",
    }).status_code == 200
    assert client.post("/api/outreach/messages", json={
        "campaign_key": "bot-outreach-panel-test",
        "recipient": address,
        "subject": "Панель рассылок",
        "body": "Тестовый черновик",
    }).status_code == 201

    response = client.get("/api/outreach/summary", headers={"X-Role": "manager"})
    assert response.status_code == 200
    summary = response.json()
    assert summary["consents"]["verified"] >= 1
    assert summary["messages"]["statuses"]["queued"] >= 1
    assert set(summary["inbound"]) == {
        "enabled",
        "receiving_ready",
        "forwarding_ready",
        "owner_destination_ready",
    }
    assert summary["safety"] == {
        "verified_consent_required": True,
        "owner_approval_required": True,
        "suppression_enforced": True,
    }
    campaign = next(row for row in summary["campaigns"]["recent"] if row["campaign_key"] == "bot-outreach-panel-test")
    assert campaign["message_count"] == 1
    assert "recipient" not in campaign
    assert client.get("/api/outreach/summary", headers={"X-Role": "viewer"}).status_code == 403


def test_telegram_outreach_panel_is_read_only_and_actionable(monkeypatch):
    from app import bot

    calls = []

    async def fake_api(method, path, **kwargs):
        calls.append((method, path, kwargs))
        return {
            "delivery_ready": True,
            "mailboxes": {"total": 2, "active": 2, "ready": 2},
            "inbound": {"enabled": 2, "receiving_ready": 2, "forwarding_ready": 2, "owner_destination_ready": True},
            "consents": {"verified": 12, "revoked": 1},
            "suppressed": 3,
            "pending_approvals": 1,
            "messages": {"total": 9, "statuses": {"queued": 4, "sent": 5}},
            "campaigns": {"total": 1, "recent": []},
            "limits": {"per_minute": 10, "per_day": 100},
        }

    class Message:
        replies = []
        async def reply_text(self, value, **kwargs): self.replies.append((value, kwargs))

    class Update:
        effective_message = Message()

    monkeypatch.setattr(bot, "api", fake_api)
    monkeypatch.setattr(bot, "allowed", lambda update: True)
    asyncio.run(bot.outreach_dashboard(Update(), None))
    assert calls == [("GET", "/api/outreach/summary", {})]
    text, kwargs = Update.effective_message.replies[0]
    assert "Подтверждённые согласия: 12" in text
    assert "Входящие ответы: готовы (2 из 2)" in text
    assert "Ожидают approval: 1" in text
    callbacks = {
        button.callback_data
        for row in kwargs["reply_markup"].inline_keyboard
        for button in row
    }
    assert callbacks == {"outreach", "outreach:campaigns", "outreach:help"}

    monkeypatch.setattr(bot, "allowed", lambda update: False)
    Update.effective_message.replies.clear()
    calls.clear()
    asyncio.run(bot.outreach_dashboard(Update(), None))
    assert calls == []
    assert Update.effective_message.replies[0][0] == "Доступ не разрешён."


def test_russian_chat_opens_outreach_panel_without_creating_a_task():
    from app.chat import understand_russian_message

    assert understand_russian_message("Покажи рассылки") == {"kind": "outreach"}
    assert understand_russian_message("Рассылки") == {"kind": "outreach"}


def test_customer_requested_campaign_records_consent_and_requires_exact_owner_approval(client, monkeypatch):
    from app.config import settings

    for field in ("smtp_host", "smtp_username", "smtp_password", "smtp_from_email"):
        monkeypatch.setattr(settings, field, "")
    payload = {
        "recipients": ["requested-b@example.com", "requested-a@example.com"],
        "consent_evidence": "12 августа 2026 клиенты передали адреса по email и попросили присылать предложения",
        "subject": "Запрошенное предложение",
        "body": "Добрый день! Направляем информацию по вашему запросу.",
    }
    draft = client.post("/api/outreach/campaigns/customer-requested/draft", json=payload)
    assert draft.status_code == 201
    result = draft.json()
    assert result["status"] == "blocked"
    assert result["recipient_count"] == 2
    assert result["approval_id"]
    assert result["idempotent_replay"] is False

    replay = client.post("/api/outreach/campaigns/customer-requested/draft", json=payload).json()
    assert replay["task_id"] == result["task_id"]
    assert replay["approval_id"] == result["approval_id"]
    assert replay["idempotent_replay"] is True

    consents = client.get("/api/outreach/consents?status=verified").json()
    addresses = {row["address"] for row in consents}
    assert {"requested-a@example.com", "requested-b@example.com"}.issubset(addresses)
    task = next(row for row in client.get("/api/tasks").json() if row["id"] == result["task_id"])
    assert task["status"] == "blocked"
    assert task["payload"]["action_kind"] == "bulk_outreach"
    assert task["payload"]["recipients"] == ["requested-a@example.com", "requested-b@example.com"]
    assert "consent_evidence" not in task["payload"]

    approved = client.post(f"/api/approvals/{result['approval_id']}/approve", json={"note": "Проверил точный preview"}).json()
    assert approved["status"] == "approved"
    assert approved["action_kind"] == "bulk_outreach"
    assert approved["execution"] == "queued"
    assert approved["task_status"] == "queued"
    task = next(row for row in client.get("/api/tasks").json() if row["id"] == result["task_id"])
    assert task["status"] == "queued"
    completed = client.post(f"/api/tasks/{result['task_id']}/run").json()
    assert completed["status"] == "blocked"
    assert completed["result"]["status"] == "credentials_required"
    assert completed["result"]["queued"] == 2
    queued = [
        row
        for row in client.get("/api/outreach/messages?status=waiting_configuration").json()
        if row["campaign_key"] == result["campaign_key"]
    ]
    assert len(queued) == 2


def test_customer_requested_campaign_keeps_attachment_bound_to_approval_and_queue(client, monkeypatch, tmp_path):
    import base64
    import hashlib

    from sqlalchemy import select

    from app.config import settings
    from app.db import SessionLocal
    from app.models import OutboundMessage

    monkeypatch.setattr(settings, "document_storage_path", str(tmp_path))
    for field in ("smtp_host", "smtp_username", "smtp_password", "smtp_from_email"):
        monkeypatch.setattr(settings, field, "")
    raw = b"%PDF-safe-customer-requested-attachment"
    payload = {
        "recipients": ["requested-attachment@example.com"],
        "consent_evidence": "13 августа 2026 клиент передал адрес и попросил присылать письма с документами",
        "subject": "Предложение с документом",
        "body": "Добрый день! Документ приложен к этому письму.",
        "attachments": [{
            "filename": "offer.pdf",
            "content_type": "application/pdf",
            "content_base64": base64.b64encode(raw).decode(),
        }],
    }
    result = client.post("/api/outreach/campaigns/customer-requested/draft", json=payload).json()
    assert result["status"] == "blocked"
    assert result["attachment_count"] == 1

    task = next(row for row in client.get("/api/tasks").json() if row["id"] == result["task_id"])
    attachment = task["payload"]["attachments"][0]
    assert attachment["filename"] == "offer.pdf"
    assert attachment["sha256"] == hashlib.sha256(raw).hexdigest()
    assert "content_base64" not in attachment
    assert Path(attachment["storage_path"]).read_bytes() == raw

    approved = client.post(
        f"/api/approvals/{result['approval_id']}/approve",
        json={"note": "Проверены адрес, текст и PDF-вложение"},
    ).json()
    assert approved["execution"] == "queued"
    completed = client.post(f"/api/tasks/{result['task_id']}/run").json()
    assert completed["result"]["waiting_configuration"] == 1
    with SessionLocal() as db:
        message = db.scalar(
            select(OutboundMessage).where(OutboundMessage.campaign_key == result["campaign_key"])
        )
        assert message is not None
        assert message.attachments[0]["sha256"] == hashlib.sha256(raw).hexdigest()


def test_customer_requested_campaign_queues_up_to_one_thousand_in_hundred_recipient_batches(client, monkeypatch):
    from app.config import settings

    for field in ("smtp_host", "smtp_username", "smtp_password", "smtp_from_email"):
        monkeypatch.setattr(settings, field, "")
    recipients = [f"batch-{index:03d}@example.com" for index in range(205)]
    payload = {
        "recipients": recipients,
        "consent_evidence": "Клиенты передали список адресов владельцу и попросили получать эту рассылку",
        "subject": "Проверка пакетной очереди 205",
        "body": "Добрый день! Это проверка защищённой пакетной очереди без реальной отправки.",
    }
    result = client.post("/api/outreach/campaigns/customer-requested/draft", json=payload).json()
    assert result["recipient_count"] == 205
    assert result["batch_size"] == 100
    assert result["batch_count"] == 3

    task = next(row for row in client.get("/api/tasks").json() if row["id"] == result["task_id"])
    assert [len(batch) for batch in task["payload"]["recipient_batches"]] == [100, 100, 5]
    assert task["status"] == "blocked"

    approved = client.post(
        f"/api/approvals/{result['approval_id']}/approve",
        json={"note": "Проверен точный список из 205 адресов, тема и текст"},
    ).json()
    assert approved["execution"] == "queued"
    completed = client.post(f"/api/tasks/{result['task_id']}/run").json()
    assert completed["status"] == "blocked"
    assert completed["result"]["status"] == "credentials_required"
    assert completed["result"]["queued"] == 205
    assert completed["result"]["waiting_configuration"] == 205
    assert completed["result"]["batch_size"] == 100
    assert completed["result"]["batch_count"] == 3
    assert [row["recipient_count"] for row in completed["result"]["batches"]] == [100, 100, 5]

    at_limit = client.post("/api/outreach/campaigns/customer-requested/draft", json={
        **payload,
        "subject": "Проверка предельной пакетной очереди 1000",
        "recipients": [f"limit-{index:04d}@example.com" for index in range(1000)],
    })
    assert at_limit.status_code == 201
    assert at_limit.json()["recipient_count"] == 1000
    assert at_limit.json()["batch_count"] == 10
    limit_task = next(row for row in client.get("/api/tasks").json() if row["id"] == at_limit.json()["task_id"])
    assert [len(batch) for batch in limit_task["payload"]["recipient_batches"]] == [100] * 10

    too_many = client.post("/api/outreach/campaigns/customer-requested/draft", json={
        **payload,
        "recipients": [f"over-{index:04d}@example.com" for index in range(1001)],
    })
    assert too_many.status_code == 422


def test_telegram_mailing_wizard_keeps_customer_data_out_of_request_analysis(monkeypatch):
    from app import bot

    calls = []

    async def fake_api(method, path, **kwargs):
        calls.append((method, path, kwargs))
        if path == "/api/telegram/control/approvals/701/card":
            return {
                "callbacks": {
                    "approve": "tc1.approve-test",
                    "reject": "tc1.reject-test",
                    "request_changes": "tc1.changes-test",
                }
            }
        return {"status": "blocked", "task_id": 700, "approval_id": 701, "recipient_count": 2, "idempotent_replay": False}

    class Message:
        text = ""
        replies = []
        async def reply_text(self, value, **kwargs): self.replies.append((value, kwargs))

    class User:
        id = 123

    class Chat:
        id = 123

    class Update:
        effective_message = Message()
        effective_user = User()
        effective_chat = Chat()

    class Context:
        args = ["one@example.com,", "two@example.com"]
        user_data = {}

    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    update, context = Update(), Context()
    asyncio.run(bot.mailing_start(update, context))
    assert context.user_data["mailing_draft"]["step"] == "evidence"

    for value, expected_step in [
        ("Клиенты передали адреса 12 августа и попросили присылать предложения", "subject"),
        ("Предложение по клинингу", "body"),
        ("Добрый день! Направляем информацию по вашему запросу.", "preview"),
    ]:
        update.effective_message.text = value
        asyncio.run(bot.natural_language(update, context))
        assert context.user_data["mailing_draft"]["step"] == expected_step
    assert calls == []
    preview_callbacks = [
        button.callback_data
        for row in update.effective_message.replies[-1][1]["reply_markup"].inline_keyboard
        for button in row
    ]
    assert "mailing:attachment" in preview_callbacks

    asyncio.run(bot.mailing_create(update, context))
    assert len(calls) == 2
    method, path, kwargs = calls[0]
    assert (method, path) == ("POST", "/api/outreach/campaigns/customer-requested/draft")
    assert kwargs["json"]["recipients"] == ["one@example.com", "two@example.com"]
    assert kwargs["json"]["attachments"] == []
    assert context.user_data.get("mailing_draft") is None
    assert "задача #700" in update.effective_message.replies[-1][0]
    markup = update.effective_message.replies[-1][1]["reply_markup"]
    assert [button.text for row in markup.inline_keyboard for button in row] == [
        "✅ Одобрить",
        "❌ Отклонить",
        "✏️ Запросить изменения",
    ]
    assert all(
        button.callback_data.startswith("tc1.")
        for row in markup.inline_keyboard
        for button in row
    )


def test_recipient_document_import_extracts_xlsx_docx_and_pdf():
    from io import BytesIO

    from docx import Document
    from openpyxl import Workbook
    from reportlab.pdfgen.canvas import Canvas

    from app.recipient_import import extract_recipient_emails

    workbook = Workbook()
    sheet = workbook.active
    for index in range(205):
        sheet.append([f"Клиент {index}", f"FILE-{index:03d}@example.com"])
    sheet.append(["Дубликат", "file-001@example.com"])
    xlsx = BytesIO()
    workbook.save(xlsx)
    extracted = extract_recipient_emails("recipients.xlsx", xlsx.getvalue())
    assert len(extracted) == 205
    assert extracted[0] == "file-000@example.com"

    structured = Workbook()
    raw = structured.active
    raw.title = "Непривязанные контакты"
    raw.append(["must-not-mail@example.com"])
    canonical = structured.create_sheet("Для импорта")
    canonical.append(["company", "email", "emails"])
    canonical.append(["УК Каноническая", "canonical@example.com", "canonical@example.com"])
    structured_xlsx = BytesIO()
    structured.save(structured_xlsx)
    assert extract_recipient_emails("structured.xlsx", structured_xlsx.getvalue()) == ["canonical@example.com"]

    document = Document()
    document.add_paragraph("Первый клиент: word-one@example.com")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "word-two@example.com"
    docx = BytesIO()
    document.save(docx)
    assert extract_recipient_emails("recipients.docx", docx.getvalue()) == [
        "word-one@example.com",
        "word-two@example.com",
    ]

    pdf = BytesIO()
    canvas = Canvas(pdf)
    canvas.drawString(50, 800, "PDF client: pdf-client@example.com")
    canvas.save()
    assert extract_recipient_emails("recipients.pdf", pdf.getvalue()) == ["pdf-client@example.com"]


def test_telegram_mailing_accepts_recipient_file_and_reports_batches(monkeypatch, tmp_path):
    from io import BytesIO

    from openpyxl import Workbook

    from app import bot

    workbook = Workbook()
    sheet = workbook.active
    for index in range(205):
        sheet.append([f"upload-{index:03d}@example.com"])
    content = BytesIO()
    workbook.save(content)

    class TelegramFile:
        async def download_to_drive(self, custom_path):
            Path(custom_path).write_bytes(content.getvalue())

    class Bot:
        async def get_file(self, file_id):
            assert file_id == "recipient-file"
            return TelegramFile()

    class Document:
        file_name = "customers.xlsx"
        file_size = len(content.getvalue())
        file_id = "recipient-file"

    class Message:
        text = ""
        document = Document()

        def __init__(self): self.replies = []
        async def reply_text(self, value, **kwargs): self.replies.append((value, kwargs))

    class User:
        id = 123

    class Update:
        def __init__(self):
            self.effective_message = Message()
            self.effective_user = User()

    class Context:
        def __init__(self):
            self.args = []
            self.user_data = {}
            self.bot = Bot()

    monkeypatch.setattr(bot, "allowed", lambda update: True)
    update, context = Update(), Context()
    asyncio.run(bot.mailing_start(update, context))
    assert context.user_data["mailing_draft"]["step"] == "document"
    assert "Excel XLSX/XLSM" in update.effective_message.replies[-1][0]

    asyncio.run(bot.proposal_document(update, context))
    draft = context.user_data["mailing_draft"]
    assert draft["step"] == "evidence"
    assert len(draft["recipients"]) == 205
    assert draft["source_filename"] == "customers.xlsx"
    assert len(draft["source_sha256"]) == 64
    assert draft["batch_size"] == 100
    assert "3 внутренних партий" in update.effective_message.replies[-1][0]


def test_telegram_mailing_attaches_business_file_to_exact_draft(monkeypatch):
    import base64

    from app import bot

    raw = b"%PDF-telegram-mail-attachment"
    calls = []

    async def fake_api(method, path, **kwargs):
        calls.append((method, path, kwargs))
        if path == "/api/telegram/control/approvals/802/card":
            return {"callbacks": {
                "approve": "tc1.approve-test",
                "reject": "tc1.reject-test",
                "request_changes": "tc1.changes-test",
            }}
        return {
            "status": "blocked",
            "task_id": 801,
            "approval_id": 802,
            "recipient_count": 1,
            "batch_count": 1,
            "batch_size": 100,
            "attachment_count": 1,
        }

    class TelegramFile:
        async def download_to_drive(self, custom_path):
            Path(custom_path).write_bytes(raw)

    class TelegramBot:
        async def get_file(self, file_id):
            assert file_id == "mail-attachment"
            return TelegramFile()

    class Document:
        file_name = "offer.pdf"
        file_size = len(raw)
        file_id = "mail-attachment"
        mime_type = "application/pdf"

    class Message:
        document = Document()

        def __init__(self):
            self.replies = []

        async def reply_text(self, value, **kwargs):
            self.replies.append((value, kwargs))

    class User:
        id = 123

    class Chat:
        id = 123

    class Update:
        effective_message = Message()
        effective_user = User()
        effective_chat = Chat()

    class Context:
        bot = TelegramBot()
        user_data = {
            "mailing_draft": {
                "step": "attachment",
                "recipients": ["attachment-wizard@example.com"],
                "consent_evidence": "Клиент запросил письма с приложенным коммерческим предложением",
                "subject": "Предложение",
                "body": "Добрый день! Направляем документ.",
            }
        }

    monkeypatch.setattr(bot, "api", fake_api)
    update, context = Update(), Context()
    assert asyncio.run(bot.mailing_document(update, context)) is True
    draft = context.user_data["mailing_draft"]
    assert draft["step"] == "preview"
    assert draft["attachments"][0]["filename"] == "offer.pdf"
    assert base64.b64decode(draft["attachments"][0]["content_base64"]) == raw
    assert "Вложение: offer.pdf" in update.effective_message.replies[-1][0]

    asyncio.run(bot.mailing_create(update, context))
    method, path, kwargs = calls[0]
    assert (method, path) == ("POST", "/api/outreach/campaigns/customer-requested/draft")
    assert kwargs["json"]["attachments"][0]["filename"] == "offer.pdf"
    assert context.user_data.get("mailing_draft") is None


def test_telegram_ambiguous_mailing_confirmation_executes_wizard(monkeypatch):
    from app import bot

    calls = []

    async def fake_api(*args, **kwargs):
        calls.append((args, kwargs))
        raise AssertionError("Ambiguous mailing must not reach request analysis or create a task")

    class Message:
        text = "запусти рассулку на client@example.com"

        def __init__(self):
            self.replies = []

        async def reply_text(self, value, **kwargs):
            self.replies.append((value, kwargs))

    class User:
        id = 123

    class Query:
        def __init__(self, data):
            self.data = data
            self.answered = False

        async def answer(self):
            self.answered = True

    class Update:
        def __init__(self):
            self.effective_message = Message()
            self.effective_user = User()
            self.callback_query = None

    class Context:
        def __init__(self):
            self.user_data = {}

    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    async def fake_authorize(update, minimum_role):
        return {"authorized": True, "role": "operator"}
    monkeypatch.setattr(bot, "_authorize_update", fake_authorize)
    update, context = Update(), Context()

    asyncio.run(bot.natural_language(update, context))
    pending = context.user_data["pending_action_confirmation"]
    assert pending["action"] == "mailing"
    assert pending["recipients"] == ["client@example.com"]
    question, kwargs = update.effective_message.replies[-1]
    assert "Вы хотели запустить рассылку?" in question
    buttons = kwargs["reply_markup"].inline_keyboard[0]
    assert [button.text for button in buttons] == ["✅ Да", "❌ Нет"]

    update.callback_query = Query(buttons[0].callback_data)
    asyncio.run(bot.callback(update, context))
    assert update.callback_query.answered is True
    assert context.user_data.get("pending_action_confirmation") is None
    assert context.user_data["mailing_draft"] == {
        "step": "evidence",
        "recipients": ["client@example.com"],
    }
    assert "Получено адресов: 1" in update.effective_message.replies[-1][0]
    assert calls == []

    asyncio.run(bot.callback(update, context))
    assert "уже устарело" in update.effective_message.replies[-1][0]
    assert calls == []


def test_telegram_ambiguous_action_waits_for_yes_then_creates_task(monkeypatch):
    from app import bot

    calls = []

    async def fake_api(method, path, **kwargs):
        calls.append((method, path, kwargs))
        if path == "/api/request-analysis":
            return {"classification": "capability_gap", "improvement_id": 44}
        if path == "/api/tasks":
            return {"id": 902, "agent_type": kwargs["json"]["agent_type"]}
        raise AssertionError(path)

    class Message:
        text = "Позвони клиенту и согласуй время"

        def __init__(self):
            self.replies = []

        async def reply_text(self, value, **kwargs):
            self.replies.append((value, kwargs))

    class User:
        id = 123

    class Query:
        def __init__(self, data): self.data = data
        async def answer(self): return None

    class Update:
        def __init__(self):
            self.effective_message = Message()
            self.effective_user = User()
            self.callback_query = None

    class Context:
        def __init__(self): self.user_data = {}

    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    async def fake_authorize(update, minimum_role):
        return {"authorized": True, "role": "operator"}
    monkeypatch.setattr(bot, "_authorize_update", fake_authorize)
    update, context = Update(), Context()

    asyncio.run(bot.natural_language(update, context))
    assert [path for _, path, _ in calls] == ["/api/request-analysis"]
    question, kwargs = update.effective_message.replies[-1]
    assert "Вы хотели создать задачу" in question
    yes = kwargs["reply_markup"].inline_keyboard[0][0]

    update.callback_query = Query(yes.callback_data)
    asyncio.run(bot.callback(update, context))
    assert [path for _, path, _ in calls] == ["/api/request-analysis", "/api/tasks"]
    assert "✅ Выполнено: создана задача #902" in update.effective_message.replies[-1][0]
    assert "улучшение #44" in update.effective_message.replies[-1][0]


def test_telegram_ambiguous_action_no_cancels_without_execution(monkeypatch):
    from app import bot

    class Message:
        text = "запусти рассулку"

        def __init__(self): self.replies = []
        async def reply_text(self, value, **kwargs): self.replies.append((value, kwargs))

    class User:
        id = 123

    class Query:
        def __init__(self, data): self.data = data
        async def answer(self): return None

    class Update:
        def __init__(self):
            self.effective_message = Message()
            self.effective_user = User()
            self.callback_query = None

    class Context:
        def __init__(self): self.user_data = {}

    monkeypatch.setattr(bot, "allowed", lambda update: True)
    async def fake_authorize(update, minimum_role):
        return {"authorized": True, "role": "operator"}
    monkeypatch.setattr(bot, "_authorize_update", fake_authorize)
    update, context = Update(), Context()
    asyncio.run(bot.natural_language(update, context))
    no = update.effective_message.replies[-1][1]["reply_markup"].inline_keyboard[0][1]
    update.callback_query = Query(no.callback_data)
    asyncio.run(bot.callback(update, context))
    assert context.user_data.get("mailing_draft") is None
    assert context.user_data.get("pending_action_confirmation") is None
    assert "Действие отменено" in update.effective_message.replies[-1][0]


def test_compose_telegram_route_is_configurable_without_a_secret():
    compose = (Path(__file__).resolve().parents[1] / "docker-compose.yml").read_text()
    route_lines = [line for line in compose.splitlines() if "api.telegram.org=" in line]
    assert len(route_lines) == 2
    assert all("api.telegram.org=${TELEGRAM_API_IP:-" in line for line in route_lines)
    assert all("TELEGRAM_BOT_TOKEN" not in line for line in route_lines)


def test_sales_agent_generates_downloadable_proposal_pdf(client, monkeypatch, tmp_path):
    from app.chat import understand_russian_message
    from app.config import settings

    monkeypatch.setattr(settings, "document_storage_path", str(tmp_path))
    monkeypatch.setattr(settings, "company_name", "CleaningAIOS")
    monkeypatch.setattr(settings, "company_legal_name", "ИП Тестовый Владелец")
    monkeypatch.setattr(settings, "company_inn", "123456789012")
    lead = client.post("/api/records", json={
        "record_type": "lead",
        "title": "Request Analyst",
        "status": "qualified",
        "source": "test",
        "data": {"name": "Request Analyst", "company": "Request Analyst", "service": "business_center", "object_area": 2500, "budget": 180000},
    }).json()
    message = "Подготовь коммерческое предложение в PDF для тестового клиента Request Analyst"
    intent = understand_russian_message(message)
    assessment = client.post("/api/request-analysis", json={"message": message, "intent": intent}).json()
    assert assessment["classification"] == "supported"
    assert assessment["improvement_id"] is None
    task = client.post("/api/tasks", json={"title": message, "agent_type": "sales", "payload": intent["payload"], "max_attempts": 1}).json()
    completed = client.post(f"/api/tasks/{task['id']}/run").json()
    assert completed["status"] == "done"
    result = completed["result"]
    assert result["status"] == "ready"
    assert result["client_record_id"] == lead["id"]
    assert result["owner_approval_required_before_sending"] is True
    assert result["sent_to_client"] is False
    downloaded = client.get(result["download_url"])
    assert downloaded.status_code == 200
    assert downloaded.headers["content-type"] == "application/pdf"
    assert downloaded.content.startswith(b"%PDF")
    assert len(downloaded.content) > 10_000
    proposals = client.get("/api/records?record_type=proposal").json()
    assert any(row["id"] == result["proposal_id"] and row["status"] == "ready" for row in proposals)
    audit = client.get("/api/audit").json()
    assert any(row["action"] == "proposal.downloaded" and row["resource_id"] == str(result["proposal_id"]) for row in audit)


def test_sales_agent_proposal_failure_is_recorded(client, monkeypatch, tmp_path):
    from app.chat import understand_russian_message
    from app.config import settings

    monkeypatch.setattr(settings, "document_storage_path", str(tmp_path))
    intent = understand_russian_message("Подготовь коммерческое предложение для клиента Отсутствует Уникально")
    task = client.post("/api/tasks", json={"title": "Missing CRM proposal", "agent_type": "sales", "payload": intent["payload"], "max_attempts": 1}).json()
    completed = client.post(f"/api/tasks/{task['id']}/run").json()
    assert completed["status"] == "failed"
    assert "не найден в CRM" in completed["result"]["error"]
    audit = client.get("/api/audit").json()
    assert any(row["action"] == "task.failed" and row["resource_id"] == str(task["id"]) for row in audit)


def test_telegram_proposal_request_returns_real_pdf(monkeypatch):
    from app import bot

    calls = []

    async def fake_api(method, path, **kwargs):
        calls.append((method, path, kwargs))
        if path == "/api/request-analysis":
            return {"classification": "supported", "improvement_id": None}
        if path == "/api/tasks":
            return {"id": 321, "agent_type": "sales", "title": "КП"}
        if path == "/api/tasks/321/run":
            return {"status": "done", "result": {"proposal_number": "KP-TEST", "download_url": "/api/proposals/77/download"}}
        raise AssertionError(path)

    async def fake_file(path):
        assert path == "/api/proposals/77/download"
        return b"%PDF-telegram-test", "proposal-test.pdf"

    class Message:
        text = "Подготовь коммерческое предложение для клиента Request Analyst"
        documents = []
        replies = []

        async def reply_document(self, **kwargs):
            self.documents.append(kwargs)

        async def reply_text(self, value, **kwargs):
            self.replies.append(value)

    class User:
        id = 123

    class Update:
        effective_message = Message()
        effective_user = User()

    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    monkeypatch.setattr(bot, "api_file", fake_file)
    asyncio.run(bot.natural_language(Update(), None))
    assert len(Update.effective_message.documents) == 1
    document = Update.effective_message.documents[0]
    assert document["filename"] == "proposal-test.pdf"
    assert document["document"].read().startswith(b"%PDF")
    assert "не отправлен клиенту" in document["caption"]
    task_payload = next(kwargs["json"] for method, path, kwargs in calls if path == "/api/tasks")
    assert task_payload["max_attempts"] == 1


def test_request_analyst_accepts_supported_request_without_improvement(client, monkeypatch):
    from app.chat import understand_russian_message
    from app.config import settings

    monkeypatch.setattr(settings, "llm_api_key", "")
    intent = understand_russian_message("Покажи текущие задачи")
    result = client.post("/api/request-analysis", json={"message": "Покажи текущие задачи", "intent": intent}).json()
    assert result["classification"] == "supported"
    assert result["fully_supported"] is True
    assert result["improvement_id"] is None


def test_request_analyst_creates_deduplicated_codex_prompt(client, monkeypatch):
    from app.chat import understand_russian_message
    from app.config import settings

    monkeypatch.setattr(settings, "llm_api_key", "")
    monkeypatch.setattr(settings, "workspace_agent_trigger_id", "")
    monkeypatch.setattr(settings, "workspace_agent_access_token", "")
    message = "Позвони клиенту и договорись о встрече"
    intent = understand_russian_message(message)
    first = client.post("/api/request-analysis", json={"message": message, "intent": intent}).json()
    second = client.post("/api/request-analysis", json={"message": message, "intent": intent}).json()
    assert first["classification"] == "capability_gap"
    assert first["improvement_id"] == second["improvement_id"]
    assert first["handoff_status"] == "credentials_required"
    queued = client.get("/api/improvements?status=queued").json()
    row = next(x for x in queued if x["id"] == first["improvement_id"])
    assert row["occurrence_count"] == 2
    assert "телефонии" in row["suggested_function"]
    assert "Required test plan" in row["codex_prompt"]
    assert row["acceptance_criteria"]
    assert row["test_plan"]


def test_request_analyst_redacts_secrets_before_storage(client, monkeypatch):
    from app.chat import understand_russian_message
    from app.config import settings

    monkeypatch.setattr(settings, "llm_api_key", "")
    secret = "123456789:ABCDEFGHIJKLMNOPQRSTUVWXYZ_abcd"
    message = f"Добавь новую функцию, token={secret}"
    intent = understand_russian_message(message)
    result = client.post("/api/request-analysis", json={"message": message, "intent": intent}).json()
    row = next(x for x in client.get("/api/improvements").json() if x["id"] == result["improvement_id"])
    assert secret not in row["request_text"]
    assert secret not in row["codex_prompt"]
    assert "[REDACTED]" in row["request_text"]


def test_request_analyst_does_not_turn_approval_policy_into_feature_gap(client, monkeypatch):
    from app.chat import understand_russian_message
    from app.config import settings

    monkeypatch.setattr(settings, "llm_api_key", "")
    message = "Оплати счет поставщику"
    result = client.post("/api/request-analysis", json={"message": message, "intent": understand_russian_message(message)}).json()
    assert result["classification"] == "approval_required"
    assert result["improvement_id"] is None


def test_workspace_agent_handoff_uses_official_trigger_contract(client, monkeypatch):
    from app.chat import understand_russian_message
    from app.config import settings
    from app import improvements as improvement_module

    monkeypatch.setattr(settings, "llm_api_key", "")
    monkeypatch.setattr(settings, "workspace_agent_trigger_id", "")
    monkeypatch.setattr(settings, "workspace_agent_access_token", "")
    message = "Позвони клиенту и согласуй время встречи"
    created = client.post("/api/request-analysis", json={"message": message, "intent": understand_russian_message(message)}).json()
    captured = {}

    class Response:
        def raise_for_status(self): return None
        def json(self): return {"conversation_url": "https://chatgpt.com/c/improvement-test", "agent_trigger_run_id": "apirun_test"}

    class Client:
        def __init__(self, *args, **kwargs): captured["client"] = kwargs
        def __enter__(self): return self
        def __exit__(self, *args): return False
        def post(self, url, headers, json): captured.update({"url": url, "headers": headers, "payload": json}); return Response()

    monkeypatch.setattr(improvement_module.httpx, "Client", Client)
    monkeypatch.setattr(settings, "workspace_agent_trigger_id", "agtch_cleaning_test")
    monkeypatch.setattr(settings, "workspace_agent_access_token", "workspace-secret")
    handed_off = client.post(f"/api/improvements/{created['improvement_id']}/handoff").json()
    assert handed_off["handoff_status"] == "queued"
    assert handed_off["workspace_run_id"] == "apirun_test"
    assert captured["url"] == "https://api.chatgpt.com/v1/workspace_agents/agtch_cleaning_test/trigger"
    assert captured["headers"]["Authorization"] == "Bearer workspace-secret"
    assert captured["headers"]["Idempotency-Key"]
    assert captured["payload"]["input"].startswith("CleaningAI OS improvement request")


def test_codex_can_update_improvement_with_test_evidence(client, monkeypatch):
    from app.chat import understand_russian_message
    from app.config import settings

    monkeypatch.setattr(settings, "llm_api_key", "")
    message = "Добавь интеграцию календаря для встреч"
    created = client.post("/api/request-analysis", json={"message": message, "intent": understand_russian_message(message)}).json()
    updated = client.patch(f"/api/improvements/{created['improvement_id']}", json={
        "status": "implemented",
        "implementation_summary": "Добавлена календарная интеграция",
        "test_evidence": [{"command": "pytest -q", "result": "passed"}],
    }).json()
    assert updated["status"] == "implemented"
    assert updated["test_evidence"][0]["result"] == "passed"


def test_request_analyst_llm_uses_strict_structured_output(monkeypatch):
    import json
    from app import llm
    from app.config import settings

    captured = {}

    class Response:
        def raise_for_status(self): return None
        def json(self):
            analysis = {
                "capability_score": 0.3,
                "reason": "Нет исполняемой функции",
                "missing_capabilities": ["proposal_generator"],
                "suggested_function": "Добавить генератор КП",
                "acceptance_criteria": ["КП создаётся из CRM"],
                "test_plan": ["Проверить PDF"],
                "should_create_improvement": True,
            }
            return {"status": "completed", "model": "test-model", "output_text": json.dumps(analysis)}

    class Client:
        def __init__(self, *args, **kwargs): captured["headers"] = kwargs["headers"]
        def __enter__(self): return self
        def __exit__(self, *args): return False
        def post(self, url, json): captured.update({"url": url, "payload": json}); return Response()

    monkeypatch.setattr(llm.httpx, "Client", Client)
    monkeypatch.setattr(settings, "llm_api_key", "test-secret")
    monkeypatch.setattr(settings, "llm_base_url", "https://api.example/v1")
    monkeypatch.setattr(settings, "llm_model", "test-model")
    result = llm.llm_advisor.analyze_request("Подготовь КП", {"kind": "task"}, {"should_create_improvement": True})
    assert result["status"] == "succeeded"
    assert result["should_create_improvement"] is True
    assert captured["payload"]["text"]["format"]["strict"] is True
    assert captured["payload"]["store"] is False


def test_orchestrator_revises_attached_proposal_with_copy_and_creative_agents(client, monkeypatch, tmp_path):
    from docx import Document
    from app.config import settings

    monkeypatch.setattr(settings, "document_storage_path", str(tmp_path))
    monkeypatch.setattr(settings, "company_name", "CleaningAIOS")
    monkeypatch.setattr(settings, "company_legal_name", "ИП Тестовый Владелец")
    monkeypatch.setattr(settings, "company_inn", "123456789012")
    source = tmp_path / "КП ИП СОКОЛОВ А С ЖК РЕЧНОЙ.docx"
    document = Document()
    document.add_paragraph("Коммерческое предложение для ЖК Речной")
    document.add_paragraph("Площадь объекта: 4 500 м²")
    document.add_paragraph("Стоимость услуг: 320 000 руб. в месяц")
    document.add_paragraph("График: ежедневно, 2 смены")
    document.save(source)

    task = client.post("/api/tasks", json={
        "title": "Обновить приложенное КП",
        "agent_type": "orchestrator",
        "priority": "high",
        "payload": {
            "action": "revise_proposal",
            "source": "telegram_document",
            "source_path": str(source),
            "source_filename": source.name,
            "request_text": "Сделай КП красивее и профессиональнее, представь мне на утверждение",
            "original_message": "Сделай КП красивее и профессиональнее, представь мне на утверждение",
        },
        "max_attempts": 1,
    }).json()
    completed = client.post(f"/api/tasks/{task['id']}/run").json()
    result = completed["result"]
    assert completed["status"] == "done"
    assert result["status"] == "ready_for_owner_review"
    assert result["owner_approval_required_before_sending"] is True
    assert result["sent_to_client"] is False
    assert {item["type"] for item in result["evidence"]} >= {
        "proposal_copy_review", "proposal_design_review", "document_export"
    }

    tasks = client.get("/api/tasks").json()
    copy_task = next(row for row in tasks if row["id"] == result["copy_task_id"])
    creative_task = next(row for row in tasks if row["id"] == result["creative_task_id"])
    assert copy_task["agent_type"] == "copywriter"
    assert copy_task["result"]["external_ai_used"] is False
    assert creative_task["agent_type"] == "creative"
    assert creative_task["result"]["preset"] == "narrative_proposal"

    docx_response = client.get(result["download_urls"]["docx"])
    pdf_response = client.get(result["download_urls"]["pdf"])
    assert docx_response.status_code == 200
    assert docx_response.content.startswith(b"PK")
    assert pdf_response.status_code == 200
    assert pdf_response.content.startswith(b"%PDF")

    approval = client.post(
        f"/api/approvals/{result['approval_id']}/approve",
        json={"note": "Утверждаю только проект, без отправки клиенту"},
    ).json()
    assert approval["status"] == "approved"
    revision = next(
        row for row in client.get("/api/records?record_type=proposal_revision").json()
        if row["id"] == result["proposal_revision_id"]
    )
    assert revision["status"] == "approved"
    assert revision["data"]["sent_to_client"] is False


def test_proposal_studio_separates_source_price_rows_and_total():
    from app.proposal_studio import _fact_lines

    text = """
    Уборка МОП (2 уборщица) 217 600
    Снабжение 14 900
    Придомовая территория (2 дворника) 232 545 Итоговая стоимость
    Менеджер клининга 35 400 500 445 руб.
    """
    assert _fact_lines(text) == [
        "Уборка МОП (2 уборщицы) — 217 600 руб.",
        "Снабжение — 14 900 руб.",
        "Придомовая территория (2 дворника) — 232 545 руб.",
        "Менеджер клининга — 35 400 руб.",
        "Итоговая стоимость — 500 445 руб.",
    ]


def test_proposal_studio_enforces_document_layout_preset(tmp_path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
    from pypdf import PdfReader

    from app.proposal_studio import _build_docx, _build_pdf

    copy = {
        "client_name": "ЖК Речной",
        "sections": {
            "opening": "Профессиональное предложение — на проверку владельцу.",
            "value": "Решение по задачам объекта.",
            "quality": "Контроль качества.",
            "launch": "Порядок запуска.",
        },
        "source_facts": ["Уборка МОП — 217 600 руб.", "Итоговая стоимость — 500 445 руб."],
        "source_contacts": {"phone": "8-995-599-60-95"},
    }
    docx_path = tmp_path / "proposal.docx"
    pdf_path = tmp_path / "proposal.pdf"
    _build_docx(docx_path, copy)
    _build_pdf(pdf_path, copy)

    document = Document(docx_path)
    section = document.sections[0]
    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)
    assert section.top_margin == section.right_margin == section.bottom_margin == section.left_margin == Inches(1)
    normal = document.styles["Normal"]
    assert normal.font.name == "Calibri"
    assert normal.font.size == Pt(11)
    assert normal.paragraph_format.space_after == Pt(8)
    assert normal.paragraph_format.line_spacing == pytest.approx(1.333, abs=0.001)
    assert normal.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    heading = document.styles["Heading 1"]
    assert heading.font.size == Pt(16)
    assert str(heading.font.color.rgb) == "2E74B5"
    assert heading.paragraph_format.space_before == Pt(18)
    assert heading.paragraph_format.space_after == Pt(10)

    for table in document.tables:
        assert table.style.name == "Table Grid"
        width = table._tbl.tblPr.first_child_found_in("w:tblW")
        assert width.get(qn("w:w")) == "9360"
        assert width.get(qn("w:type")) == "dxa"
        assert sum(int(node.get(qn("w:w"))) for node in table._tbl.tblGrid) == 9360

    assert len(PdfReader(str(pdf_path)).pages) == 2
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(str(pdf_path)).pages)
    assert "—" not in pdf_text
    assert "ПРОЕКТ · ТРЕБУЕТ ПРОВЕРКИ И УТВЕРЖДЕНИЯ ВЛАДЕЛЬЦА" in pdf_text


def test_incomplete_telegram_task_creates_deduplicated_improvement_and_ceo_report(client, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "workspace_agent_trigger_id", "")
    monkeypatch.setattr(settings, "workspace_agent_access_token", "")
    task = client.post("/api/tasks", json={
        "title": "Обработать крупное приложенное КП",
        "agent_type": "orchestrator",
        "payload": {
            "action": "revise_proposal",
            "source": "telegram_document",
            "document_status": "credentials_required",
            "source_filename": "large-proposal.docx",
        },
        "max_attempts": 1,
    }).json()
    first = client.post(f"/api/tasks/{task['id']}/run").json()
    assert first["status"] == "blocked"
    assert first["result"]["improvement_id"]
    assert first["result"]["ceo_incident_task_id"]
    assert first["result"]["handoff_status"] == "credentials_required"
    assert first["result"]["responsible_party"] == "owner_configuration"

    incident = next(
        row for row in client.get("/api/tasks").json()
        if row["id"] == first["result"]["ceo_incident_task_id"]
    )
    assert incident["agent_type"] == "ceo"
    assert incident["result"]["report_kind"] == "agent_incident"
    assert incident["result"]["source_task_id"] == task["id"]
    runs = client.get("/api/agent-runs", headers={"X-Role": "manager"}).json()
    source_run = next(row for row in runs if row["task_id"] == task["id"])
    assert source_run["status"] == "incomplete"


def test_telegram_large_proposal_is_not_silently_ignored(monkeypatch):
    from app import bot
    from app.config import settings

    calls = []

    async def fake_api(method, path, **kwargs):
        calls.append((method, path, kwargs))
        if path == "/api/request-analysis":
            return {"classification": "supported"}
        if path == "/api/tasks":
            return {"id": 701}
        if path == "/api/tasks/701/run":
            return {"id": 701, "status": "blocked", "result": {"improvement_id": 81, "ceo_incident_task_id": 702}}
        raise AssertionError(path)

    class Document:
        file_name = "КП ЖК Речной.docx"
        file_size = 29_800_000
        file_id = "telegram-file-id"

    class Message:
        document = Document()
        caption = "Сделай коммерческое предложение более красивым и профессиональным и представь на утверждение"

        def __init__(self):
            self.replies = []

        async def reply_text(self, value, **kwargs):
            self.replies.append(value)

    class User:
        id = 123

    class Update:
        effective_message = Message()
        effective_user = User()

    monkeypatch.setattr(settings, "telegram_bot_api_base_url", "")
    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    asyncio.run(bot.proposal_document(Update(), None))
    assert "зарегистрированы как задача #701" in Update.effective_message.replies[-1]
    assert "улучшение #81" in Update.effective_message.replies[-1].lower()
    task_payload = next(kwargs["json"] for method, path, kwargs in calls if path == "/api/tasks")
    assert task_payload["payload"]["document_status"] == "credentials_required"
    assert task_payload["payload"]["source"] == "telegram_document"


def test_telegram_small_proposal_returns_docx_and_pdf_for_owner_review(monkeypatch, tmp_path):
    from app import bot
    from app.config import settings

    calls = []

    async def fake_api(method, path, **kwargs):
        calls.append((method, path, kwargs))
        if path == "/api/request-analysis":
            return {"classification": "supported"}
        if path == "/api/tasks":
            return {"id": 801}
        if path == "/api/tasks/801/run":
            return {
                "id": 801,
                "status": "done",
                "result": {
                    "status": "ready_for_owner_review",
                    "proposal_number": "KPR-TEST",
                    "approval_id": 91,
                    "download_urls": {"docx": "/revision.docx", "pdf": "/revision.pdf"},
                },
            }
        if path == "/api/telegram/control/approvals/91/card":
            return {
                "callbacks": {
                    "approve": "tc1.approve-test",
                    "reject": "tc1.reject-test",
                    "request_changes": "tc1.changes-test",
                }
            }
        raise AssertionError(path)

    async def fake_file(path):
        return (b"PK-test", "revision.docx") if path.endswith("docx") else (b"%PDF-test", "revision.pdf")

    class TelegramFile:
        async def download_to_drive(self, custom_path):
            Path(custom_path).write_bytes(b"source-docx")

    class TelegramBot:
        async def get_file(self, file_id):
            assert file_id == "file-id"
            return TelegramFile()

    class Context:
        bot = TelegramBot()

    class Document:
        file_name = "КП ЖК Речной.docx"
        file_size = 100
        file_id = "file-id"

    class Message:
        document = Document()
        caption = "Сделай это КП профессиональнее и представь мне на утверждение"

        def __init__(self):
            self.replies = []
            self.documents = []

        async def reply_text(self, value, **kwargs):
            self.replies.append(value)

        async def reply_document(self, **kwargs):
            self.documents.append(kwargs)

    class User:
        id = 123

    class Chat:
        id = 123

    class Update:
        effective_message = Message()
        effective_user = User()
        effective_chat = Chat()

    monkeypatch.setattr(settings, "document_storage_path", str(tmp_path))
    monkeypatch.setattr(settings, "telegram_bot_api_base_url", "")
    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    monkeypatch.setattr(bot, "api_file", fake_file)
    asyncio.run(bot.proposal_document(Update(), Context()))
    assert [item["filename"] for item in Update.effective_message.documents] == ["revision.docx", "revision.pdf"]
    assert "Клиенту ничего не отправлено" in Update.effective_message.documents[-1]["caption"]
    assert Update.effective_message.documents[-1]["reply_markup"] is not None


def test_public_website_lead_enters_crm_inbox_and_hot_queue(client, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "hot_lead_score", 70)
    payload = {
        "name": "Анна Петрова",
        "company": "УК Публичный тест",
        "phone": "+7 999 123-45-67",
        "email": "public-hot@example.com",
        "service": "business_center",
        "object_area": 5000,
        "budget": 400000,
        "urgency": "today",
        "message": "Нужен быстрый запуск ежедневной уборки",
        "consent": True,
        "utm_source": "yandex",
        "utm_medium": "cpc",
        "utm_campaign": "public-hot-test",
    }
    response = client.post("/api/public/leads", json=payload)
    assert response.status_code == 201
    result = response.json()
    assert result["status"] == "qualified"
    assert result["owner_notification"] in {"queued", "credentials_required"}
    leads = client.get("/api/records?record_type=lead").json()
    lead = next(row for row in leads if row["id"] == result["lead_id"])
    assert lead["source"] == "yandex"
    assert lead["data"]["utm_campaign"] == "public-hot-test"
    inbox = client.get("/api/inbox?channel=web").json()
    assert any(row["record_id"] == result["lead_id"] for row in inbox)
    tasks = client.get("/api/tasks").json()
    assert any(row["agent_type"] == "sales" and row["payload"].get("record_id") == result["lead_id"] for row in tasks)


def test_public_lead_requires_contact_and_consent(client):
    base = {"name": "Тест", "service": "commercial", "urgency": "month", "consent": True}
    assert client.post("/api/public/leads", json=base).status_code == 422
    assert client.post("/api/public/leads", json={**base, "phone": "+7 999 000-00-00", "consent": False}).status_code == 422


def test_published_website_news_and_media_workflow(client):
    content = client.post("/api/marketing/content", json={"channel": "website", "title": "Запустили новый стандарт", "body": "Проверяем качество по единому регламенту.", "status": "draft"}).json()
    published = client.patch(f"/api/marketing/content/{content['id']}", json={"status": "published", "metrics": {"cover_url": "/static/cleaning-hero.png"}})
    assert published.status_code == 200
    asset = client.post("/api/marketing/media-assets", json={"kind": "image", "title": "Визуал новости", "prompt": "Минималистичный чистый интерьер"}).json()
    assert asset["provider"] == "codex_imagegen_workflow"
    completed = client.patch(f"/api/marketing/media-assets/{asset['id']}", headers={"X-Role": "manager"}, json={"status": "published", "public_url": "/static/cleaning-hero.png", "alt_text": "Чистый интерьер"})
    assert completed.status_code == 200
    site = client.get("/api/public/site").json()
    assert any(row["id"] == content["id"] for row in site["news"])
    assert any(row["id"] == asset["id"] for row in site["media"])


def test_marketing_experiment_attribution_and_manual_activation(client):
    experiment = client.post("/api/marketing/experiments", headers={"X-Role": "manager"}, json={
        "title": "Спрос на уборку БЦ",
        "channel": "yandex_direct",
        "hypothesis": "Уточнение SLA увеличит число квалифицированных лидов",
        "audience": "Управляющие бизнес-центров",
        "offer": "Аудит объекта до расчёта",
        "primary_metric": "qualified_leads",
        "budget_limit": 0,
        "utm_campaign": "experiment-attribution-test",
    }).json()
    waiting = client.post(f"/api/marketing/experiments/{experiment['id']}/launch", headers={"X-Role": "manager"}, json={}).json()
    assert waiting["status"] == "approved_waiting_manual_activation"
    started = client.post(f"/api/marketing/experiments/{experiment['id']}/launch", headers={"X-Role": "manager"}, json={"external_campaign_id": "manual-ya-123"}).json()
    assert started["status"] == "running"
    assert started["automatic_spend"] is False
    client.post("/api/public/leads", json={"name": "Иван Тест", "phone": "+7 999 111-22-33", "service": "business_center", "urgency": "week", "consent": True, "utm_campaign": "experiment-attribution-test"})
    analytics = client.get(f"/api/marketing/experiments/{experiment['id']}/analytics").json()
    assert analytics["leads"] == 1


def test_marketing_invoice_routes_to_owner_without_payment(client):
    provider = client.post("/api/marketing/providers", json={"name": "Рекламное агентство Тест", "platform": "agency", "contact": "manager@example.com"}).json()
    requisites = client.post("/api/company/requisites", json={
        "profile_name": "Основные тестовые",
        "legal_name": "ООО Тест Клининг",
        "inn": "7701234567",
        "kpp": "770101001",
        "settlement_account": "40702810000000000001",
        "currency": "RUR",
        "bank_name": "Тестовый банк",
        "bank_inn": "7701234567",
        "bank_address": "Москва, тестовый адрес банка",
        "bic": "044525001",
        "correspondent_account": "30101810000000000001",
    }).json()
    assert requisites["currency"] == "RUB"
    assert requisites["bank_inn"].endswith("4567")
    assert requisites["bank_address"] == "Москва, тестовый адрес банка"
    invoice = client.post("/api/marketing/invoices", headers={"X-Role": "manager"}, json={
        "provider_id": provider["id"],
        "requisites_profile_id": requisites["id"],
        "invoice_number": "MKT-001",
        "amount": 25000,
        "description": "Тест рекламной гипотезы",
    }).json()
    assert invoice["status"] == "pending_approval"
    assert invoice["automatic_payment"] is False
    assert invoice["telegram_notification"] in {"queued", "credentials_required"}
    decision = client.post(f"/api/approvals/{invoice['approval_id']}/approve", json={"note": "Проверено"}).json()
    assert decision["execution"] == "not_executed"
    saved = next(row for row in client.get("/api/marketing/invoices", headers={"X-Role": "manager"}).json() if row["id"] == invoice["id"])
    assert saved["status"] == "approved_for_manual_payment"
    assert saved["automatic_payment"] is False


def test_ai_provider_router_has_least_privilege_policy(client):
    response = client.get("/api/ai/providers", headers={"X-Role": "manager"})
    assert response.status_code == 200
    result = response.json()
    assert result["blanket_access"] is False
    assert result["policy"] == "least_privilege"
    assert all(row["forbidden"] for row in result["providers"])


def test_telegram_owner_notification_uses_existing_approval_buttons(monkeypatch):
    from app import notifications
    from app.config import settings
    from app.models import ApprovalRequest, OwnerNotification
    from app.telegram_control import parse_callback_token

    captured = {}

    class Response:
        def raise_for_status(self): return None

    class Client:
        def __init__(self, *args, **kwargs): captured["client"] = kwargs
        def __enter__(self): return self
        def __exit__(self, *args): return False
        def post(self, url, json): captured.update({"url": url, "payload": json}); return Response()

    monkeypatch.setattr(notifications.httpx, "Client", Client)
    monkeypatch.setattr(settings, "telegram_bot_token", "123456:test-token-value")
    monkeypatch.setattr(settings, "owner_telegram_id", "999")
    monkeypatch.setattr(settings, "telegram_callback_secret", "test-callback-secret")
    approval = ApprovalRequest(
        id=42,
        action_kind="financial",
        resource_type="marketing_invoice",
        resource_id="42",
        rationale="Проверить счёт",
        status="pending",
        decision_version=1,
        payload={},
    )
    class FakeDb:
        def get(self, model, identity):
            assert (model, identity) == (ApprovalRequest, 42)
            return approval
    row = OwnerNotification(
        idempotency_key="telegram-contract-test",
        channel="telegram",
        recipient="999",
        subject="Счёт на рекламу",
        body="Одобрение не выполняет оплату",
        data={"approval_id": approval.id},
    )
    notifications._send_telegram(FakeDb(), row)
    assert captured["url"].endswith("/sendMessage")
    assert captured["payload"]["chat_id"] == "999"
    buttons = [
        item
        for group in captured["payload"]["reply_markup"]["inline_keyboard"]
        for item in group
    ]
    assert [item["text"] for item in buttons] == [
        "✅ Одобрить",
        "❌ Отклонить",
        "✏️ Запросить изменения",
    ]
    assert {parse_callback_token(item["callback_data"])["action"] for item in buttons} == {
        "approve",
        "reject",
        "request_changes",
    }
    assert all(len(item["callback_data"].encode()) <= 64 for item in buttons)


def test_telegram_social_approval_sends_visual_album_before_buttons(monkeypatch):
    from app import notifications
    from app.config import settings
    from app.models import ApprovalRequest, OwnerNotification
    from app.telegram_control import parse_callback_token

    captured = []

    class Response:
        def raise_for_status(self): return None

    class Client:
        def __init__(self, *args, **kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *args): return False
        def post(self, url, json):
            captured.append((url, json))
            return Response()

    monkeypatch.setattr(notifications.httpx, "Client", Client)
    monkeypatch.setattr(settings, "telegram_bot_token", "123456:test-token-value")
    monkeypatch.setattr(settings, "telegram_callback_secret", "test-callback-secret")
    approval = ApprovalRequest(
        id=43,
        action_kind="social_publication",
        resource_type="social_content_batch",
        resource_id="43",
        rationale="Проверить публикации",
        status="pending",
        decision_version=1,
        payload={},
    )
    class FakeDb:
        def get(self, model, identity):
            assert (model, identity) == (ApprovalRequest, 43)
            return approval
    row = OwnerNotification(
        idempotency_key="telegram-social-preview-test",
        channel="telegram",
        recipient="999",
        subject="Визуальное согласование",
        body="Проверьте макеты",
        data={
            "approval_id": approval.id,
            "preview_posts": [{
                "channel": "vk",
                "scheduled_at": "2031-06-18T07:00:00",
                "body": "Точный текст публикации",
                "image_url": "https://cleaning.example/static/post.png",
            }],
        },
    )
    notifications._send_telegram(FakeDb(), row)
    assert captured[0][0].endswith("/sendMediaGroup")
    assert captured[0][1]["media"][0]["media"] == "https://cleaning.example/static/post.png"
    assert "Точный текст публикации" in captured[0][1]["media"][0]["caption"]
    assert captured[1][0].endswith("/sendMessage")
    callback_data = captured[1][1]["reply_markup"]["inline_keyboard"][0][0]["callback_data"]
    assert parse_callback_token(callback_data)["action"] == "approve"


def test_management_company_import_preserves_provenance_and_requires_consent(client):
    import base64

    csv = (
        "company,region,email,phone,website,inn\n"
        "УК Север Тест,Санкт-Петербург,office@uk-sever.example,+7 812 123-45-67,https://uk-sever.example,7812345678\n"
        "УК вне региона,Москва,other@example.com,,,\n"
    ).encode()
    response = client.post("/api/research/management-companies/import", json={
        "filename": "management-companies.csv",
        "content_base64": base64.b64encode(csv).decode(),
        "source_kind": "gis_housing",
        "source_url": "https://dom.gosuslugi.ru/",
    })
    assert response.status_code == 201
    result = response.json()
    assert result["created"] == 1
    assert result["skipped"] == 1
    record = next(row for row in client.get("/api/records?record_type=management_company").json() if row["title"] == "УК Север Тест")
    assert record["data"]["marketing_consent_status"] == "unknown"
    assert record["data"]["provenance"][0]["source_url"] == "https://dom.gosuslugi.ru/"
    assert record["data"]["provenance"][0]["source_filename"] == "management-companies.csv"
    assert len(record["data"]["provenance"][0]["source_sha256"]) == 64

    blocked = client.post("/api/outreach/campaigns/launch", headers={"X-Role": "manager"}, json={
        "campaign_key": "uk-no-consent-test",
        "recipients": ["office@uk-sever.example"],
        "subject": "Клининг",
        "body": "Предложение",
    })
    assert blocked.status_code == 422
    assert blocked.json()["detail"]["count"] == 1


def test_management_company_xlsx_uses_canonical_import_sheet_and_never_trusts_consent(client):
    import base64
    from io import BytesIO

    from openpyxl import Workbook

    workbook = Workbook()
    summary = workbook.active
    summary.title = "Сводка"
    summary.append(["Показатель", "Значение"])
    summary.append(["Уникальных email", 3])
    source = workbook.create_sheet("Непривязанные контакты")
    source.append(["Email"])
    source.append(["must-not-import@example.com"])
    import_sheet = workbook.create_sheet("Для импорта")
    import_sheet.append(["Данные для импорта CleaningAIOS"])
    import_sheet.append(["Согласие остаётся unknown, импорт не запускает рассылку"])
    import_sheet.append([])
    import_sheet.append([])
    import_sheet.append([
        "organization_type",
        "company",
        "region",
        "email",
        "marketing_consent_status",
        "source_refs",
    ])
    import_sheet.append(["УК", "УК Канонический XLSX", "Санкт-Петербург", "canonical@example.com", "confirmed", "owner:xlsx:1"])
    content = BytesIO()
    workbook.save(content)

    response = client.post("/api/research/management-companies/import", json={
        "filename": "canonical-management-companies.xlsx",
        "content_base64": base64.b64encode(content.getvalue()).decode(),
        "source_kind": "manual_public_export",
        "source_url": "owner-upload://canonical-management-companies",
    })
    assert response.status_code == 201
    result = response.json()
    assert result["total_rows"] == 1
    assert len(result["source_sha256"]) == 64
    records = client.get("/api/records?record_type=management_company").json()
    record = next(row for row in records if row["title"] == "УК Канонический XLSX")
    assert record["data"]["emails"] == ["canonical@example.com"]
    assert record["data"]["marketing_consent_status"] == "unknown"
    assert record["data"]["provenance"][0]["source_filename"] == "canonical-management-companies.xlsx"
    assert all("must-not-import@example.com" not in (row["data"].get("emails") or []) for row in records)


def test_management_company_import_keeps_structured_contacts_without_trusting_consent(client):
    import base64

    csv = (
        "organization_type,company,region,scope_status,address,phone,phones,email,emails,website,"
        "candidate_website,vk_url,contact_person,managed_objects,managed_area_m2,marketing_consent_status,"
        "internet_verification_status,source_refs\n"
        "УК,УК Структура Альфа,Санкт-Петербург,in_scope,Невский проспект 1,+7 812 100-00-01,"
        "+7 812 100-00-01 | +7 812 100-00-02,shared-uk@example.com,"
        "shared-uk@example.com | director-alpha@example.com,,https://alpha.example,https://vk.com/alpha,"
        "Иван Иванов,4,12345.5,confirmed,not_checked,docx:table:1:row:2\n"
        "ТСЖ,ТСЖ Структура Бета,Ленинградская область,in_scope,,,+7 812 200-00-01,"
        "shared-uk@example.com,shared-uk@example.com,,,,,1,,confirmed,not_checked,docx:table:2:row:4\n"
        "ТСЖ,ТСЖ Структура Без Email,Санкт-Петербург,in_scope,,+7 812 300-00-01,,,,,,,,,,unknown,"
        "not_checked,docx:table:3:row:6\n"
    ).encode()
    response = client.post("/api/research/management-companies/import", json={
        "filename": "structured-management-companies.csv",
        "content_base64": base64.b64encode(csv).decode(),
        "source_kind": "manual_public_export",
        "source_url": "owner-upload://structured-management-companies",
    })
    assert response.status_code == 201
    assert response.json()["created"] == 3
    records = {
        row["title"]: row
        for row in client.get("/api/records?record_type=management_company").json()
        if row["title"].startswith(("УК Структура", "ТСЖ Структура"))
    }
    assert set(records) == {"УК Структура Альфа", "ТСЖ Структура Бета", "ТСЖ Структура Без Email"}
    alpha = records["УК Структура Альфа"]["data"]
    assert alpha["emails"] == ["director-alpha@example.com", "shared-uk@example.com"]
    assert alpha["phones"] == ["+78121000001", "+78121000002"]
    assert alpha["organization_type"] == "УК"
    assert alpha["managed_objects"] == 4
    assert alpha["candidate_website"] == "https://alpha.example"
    assert alpha["source_refs"] == ["docx:table:1:row:2"]
    assert alpha["marketing_consent_status"] == "unknown"
    assert records["ТСЖ Структура Бета"]["id"] != records["УК Структура Альфа"]["id"]


def test_research_and_sales_agents_audit_management_company_base_without_sending(client):
    from app.db import SessionLocal
    from app.models import OutboundMessage
    from sqlalchemy import func, select

    with SessionLocal() as db:
        messages_before = db.scalar(select(func.count(OutboundMessage.id))) or 0

    research = client.post("/api/tasks", json={
        "title": "Проверить полноту контактов УК и ТСЖ",
        "agent_type": "research",
        "payload": {
            "collection": "management_company_contacts",
            "batch_limit": 5,
            "enrich_verified_websites": False,
        },
        "max_attempts": 1,
    }).json()
    research_done = client.post(f"/api/tasks/{research['id']}/run").json()
    assert research_done["status"] == "done"
    coverage = research_done["result"]
    assert coverage["organizations_missing_email"] >= 1
    assert coverage["search_adapter_required"] is True
    assert coverage["send_eligible_addresses"] == 0
    assert coverage["evidence"][0]["type"] == "management_company_contact_coverage"

    sales = client.post("/api/tasks", json={
        "title": "Сегментировать базу УК и ТСЖ без запуска рассылки",
        "agent_type": "sales",
        "payload": {"action": "prepare_management_company_outreach"},
        "max_attempts": 1,
    }).json()
    sales_done = client.post(f"/api/tasks/{sales['id']}/run").json()
    assert sales_done["status"] == "done"
    summary = sales_done["result"]
    assert summary["segments_by_type"]["УК"] >= 1
    assert summary["segments_by_type"]["ТСЖ"] >= 2
    assert summary["send_eligible_addresses"] == 0
    assert summary["messages_queued"] == 0
    assert summary["owner_approval_required_before_campaign"] is True
    with SessionLocal() as db:
        assert (db.scalar(select(func.count(OutboundMessage.id))) or 0) == messages_before


def test_management_company_internet_discovery_reports_real_adapter_blocker(client):
    task = client.post("/api/tasks", json={
        "title": "Проверить границу интернет-поиска УК",
        "agent_type": "research",
        "payload": {"collection": "management_company_internet_discovery"},
        "max_attempts": 1,
    }).json()
    completed = client.post(f"/api/tasks/{task['id']}/run").json()
    assert completed["status"] == "done"
    result = completed["result"]
    assert result["status"] == "adapter_required"
    assert result["configured"] is False
    assert result["credentials_required"] == ["YANDEX_SEARCH_API_KEY", "YANDEX_CLOUD_FOLDER_ID"]
    assert result["evidence"] == []


def test_campaign_attachment_balances_only_consented_recipients(client, monkeypatch, tmp_path):
    import base64

    from app.config import settings
    from app.db import SessionLocal
    from app.models import OutboundMessage
    from sqlalchemy import select

    monkeypatch.setattr(settings, "document_storage_path", str(tmp_path))
    monkeypatch.setenv("SMTP_POOL_A", "safe-test-secret-a")
    monkeypatch.setenv("SMTP_POOL_B", "safe-test-secret-b")
    first = client.post("/api/outreach/mailboxes", json={
        "name": "Gmail pool A",
        "address": "pool-a@example.com",
        "smtp_host": "smtp.gmail.com",
        "secret_ref": "SMTP_POOL_A",
        "per_minute": 2,
        "per_day": 50,
    }).json()
    second = client.post("/api/outreach/mailboxes", json={
        "name": "Gmail pool B",
        "address": "pool-b@example.com",
        "smtp_host": "smtp.gmail.com",
        "secret_ref": "SMTP_POOL_B",
        "per_minute": 2,
        "per_day": 50,
    }).json()
    recipients = ["consented-a@example.com", "consented-b@example.com"]
    for address in recipients:
        assert client.put("/api/outreach/consents", json={
            "address": address,
            "source_url": f"https://consent.example.test/{address}",
            "evidence": f"Opt-in evidence recorded for {address}",
        }).status_code == 200
    launch = {
        "campaign_key": "balanced-attachment-test",
        "recipients": recipients,
        "subject": "Предложение",
        "body": "Добрый день",
        "attachments": [{
            "filename": "offer.pdf",
            "content_type": "application/pdf",
            "content_base64": base64.b64encode(b"%PDF-safe-test").decode(),
        }],
        "auto_balance_mailboxes": True,
    }
    waiting = client.post("/api/outreach/campaigns/launch", headers={"X-Role": "manager"}, json=launch).json()
    assert waiting["status"] == "waiting_approval"
    client.post(f"/api/approvals/{waiting['approval_id']}/approve", json={"note": "Approved exact recipients and attachment"})
    launch["approval_id"] = waiting["approval_id"]
    queued = client.post("/api/outreach/campaigns/launch", headers={"X-Role": "manager"}, json=launch).json()
    assert queued["queued"] == 2
    assert queued["mailbox_distribution"] == {str(first["id"]): 1, str(second["id"]): 1}
    with SessionLocal() as db:
        messages = db.scalars(select(OutboundMessage).where(OutboundMessage.campaign_key == launch["campaign_key"])).all()
        assert all(message.attachments[0]["storage_path"].startswith(str(tmp_path)) for message in messages)
        assert all("content_base64" not in message.attachments[0] for message in messages)


def test_inbound_mail_is_deduplicated_and_forwarded_through_originating_mailbox(client, monkeypatch):
    from email.message import EmailMessage

    from app import inbound_mail, notifications
    from app.config import settings
    from app.db import SessionLocal
    from app.models import InboxMessage, OwnerNotification, SenderMailbox
    from sqlalchemy import func, select, update

    raw = EmailMessage()
    raw["From"] = "director@uk.example"
    raw["To"] = "pool-inbound@example.com"
    raw["Subject"] = "Re: Клининг"
    raw["Message-ID"] = "<inbound-unique@example.com>"
    raw.set_content("Интересно, пришлите расчёт.")
    message_bytes = raw.as_bytes()

    class IMAP:
        def __init__(self, *args, **kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *args): return False
        def login(self, username, password): return "OK", []
        def select(self, mailbox, readonly=True): return "OK", [b"1"]
        def uid(self, command, *args):
            if command == "search": return "OK", [b"7"]
            return "OK", [(b"7 (RFC822 {100})", message_bytes)]
        def logout(self): return "BYE", []

    forwarded = []

    class SMTP:
        def __init__(self, *args, **kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *args): return False
        def starttls(self): return None
        def login(self, username, password): return None
        def send_message(self, message): forwarded.append(message)

    monkeypatch.setenv("IMAP_POOL_TEST", "test-app-password")
    monkeypatch.setenv("SMTP_POOL_TEST", "test-app-password")
    monkeypatch.setattr(inbound_mail.imaplib, "IMAP4_SSL", IMAP)
    monkeypatch.setattr(notifications.smtplib, "SMTP", SMTP)
    monkeypatch.setattr(settings, "owner_notification_email", "owner@example.com")
    for field in ("smtp_host", "smtp_username", "smtp_password", "smtp_from_email"):
        monkeypatch.setattr(settings, field, "")
    with SessionLocal() as db:
        mailbox = SenderMailbox(
            name="Inbound test",
            address="pool-inbound@example.com",
            smtp_host="smtp.gmail.com",
            username="pool-inbound@example.com",
            secret_ref="SMTP_POOL_TEST",
            imap_host="imap.gmail.com",
            imap_secret_ref="IMAP_POOL_TEST",
            inbound_enabled=True,
        )
        db.add(mailbox)
        db.commit(); db.refresh(mailbox)
        first = inbound_mail.collect_mailbox_replies(db, mailbox)
        db.commit()
        assert first["received"] == 1
        mailbox.last_imap_uid = 0
        db.commit()
        second = inbound_mail.collect_mailbox_replies(db, mailbox)
        db.commit()
        assert second["duplicates"] == 1
        assert db.scalar(select(func.count()).select_from(InboxMessage).where(InboxMessage.external_id == "<inbound-unique@example.com>")) == 1
        notification = db.scalar(select(OwnerNotification).where(OwnerNotification.idempotency_key == f"inbound-email:{mailbox.id}:7"))
        assert notification is not None
        assert notification.status == "queued"
        assert notification.data["reply_to"] == "director@uk.example"
        inbox = db.get(InboxMessage, int(notification.resource_id))
        assert inbox.data["forwarded_to_owner"] is False

        db.execute(update(OwnerNotification).where(OwnerNotification.id != notification.id).values(status="sent"))
        db.commit()
        assert notifications.send_next_owner_notification(db) is True
        db.refresh(notification)
        db.refresh(inbox)
        assert notification.status == "sent"
        assert inbox.data["forwarded_to_owner"] is True
        assert inbox.data["owner_notification_id"] == notification.id

        monkeypatch.setattr(settings, "owner_notification_email", mailbox.address)
        loop_guard = notifications.queue_owner_notification(
            db,
            idempotency_key=f"inbound-loop-guard:{mailbox.id}",
            channel="email",
            resource_type="inbox_message",
            resource_id="999999",
            subject="Loop guard",
            body="Must not be sent back to the monitored inbox",
            data={"mailbox_id": mailbox.id},
        )
        assert loop_guard.status == "waiting_configuration"

    assert len(forwarded) == 1
    assert forwarded[0]["From"] == "pool-inbound@example.com"
    assert forwarded[0]["To"] == "owner@example.com"
    assert forwarded[0]["Reply-To"] == "director@uk.example"


def test_marketing_agent_creates_two_posts_per_social_channel_for_approval(client):
    task = client.post("/api/tasks", json={
        "title": "Daily social plan deterministic test",
        "agent_type": "marketing",
        "payload": {"action": "prepare_daily_social_plan", "day": "2031-06-18T06:00:00+00:00"},
        "max_attempts": 1,
    }).json()
    completed = client.post(f"/api/tasks/{task['id']}/run").json()
    assert completed["status"] == "done"
    result = completed["result"]
    assert result["created"] == 8
    assert result["approval_id"] is None
    assert len(result["media_asset_ids"]) == 2
    items = [row for row in client.get("/api/marketing/content?status=visual_pending").json() if row["id"] in result["content_item_ids"]]
    assert len(items) == 8
    assert {row["channel"] for row in items} == {"telegram", "vk", "odnoklassniki", "instagram"}
    assert all(sum(row["channel"] == channel for row in items) == 2 for channel in {"telegram", "vk", "odnoklassniki", "instagram"})
    assets = [row for row in client.get("/api/marketing/media-assets?status=queued").json() if row["id"] in result["media_asset_ids"]]
    assert len(assets) == 2
    assert all(row["provider"] == "openai_images" for row in assets)
    assert all(row["metadata"]["visual_review_required"] is True for row in assets)

    first = client.patch(f"/api/marketing/media-assets/{assets[0]['id']}", json={
        "provider": "imagegen",
        "public_url": "/static/og.png",
        "alt_text": "Проверенный визуал первого поста",
        "status": "ready",
        "metadata": {"visually_reviewed": True, "sha256": "a" * 64},
    }).json()
    assert first["social_preview"]["status"] == "visuals_pending"
    assert first["social_preview"]["approval_id"] is None
    second = client.patch(f"/api/marketing/media-assets/{assets[1]['id']}", json={
        "provider": "imagegen",
        "public_url": "/static/cleaning-hero.png",
        "alt_text": "Проверенный визуал второго поста",
        "status": "ready",
        "metadata": {"visually_reviewed": True, "sha256": "b" * 64},
    }).json()
    approval_id = second["social_preview"]["approval_id"]
    assert approval_id
    preview = client.get(f"/api/marketing/social-batches/{result['batch_id']}/preview", headers={"X-Role": "manager"}).json()
    assert preview["all_visuals_ready"] is True
    assert len(preview["posts"]) == 8
    assert all(post["image_url"] and post["body"] for post in preview["posts"])
    assert len({post["visual_asset_id"] for post in preview["posts"]}) == 2

    first_post = preview["posts"][0]
    client.patch(f"/api/marketing/content/{first_post['content_item_id']}", json={"body": first_post["body"] + " изменено"})
    stale = client.post(f"/api/approvals/{approval_id}/approve", json={"note": "Must not approve changed content"})
    assert stale.status_code == 409
    assert "changed after the preview" in stale.json()["detail"]
    client.patch(f"/api/marketing/content/{first_post['content_item_id']}", json={"body": first_post["body"]})
    approved = client.post(f"/api/approvals/{approval_id}/approve", json={"note": "Reviewed every image and caption"})
    assert approved.status_code == 200
    scheduled = client.get("/api/marketing/content?status=scheduled").json()
    scheduled_items = [row for row in scheduled if row["id"] in result["content_item_ids"]]
    assert len(scheduled_items) == 6
    assert {row["channel"] for row in scheduled_items} == {"telegram", "vk", "odnoklassniki"}
    instagram_drafts = [row for row in client.get("/api/marketing/content?status=approval").json() if row["id"] in result["content_item_ids"]]
    assert len(instagram_drafts) == 2
    assert all(row["metrics"]["publication_status"] == "legal_review_required" for row in instagram_drafts)


def test_telegram_document_creates_guarded_outreach_draft(monkeypatch, tmp_path):
    from app import bot
    from app.config import settings

    calls = []

    async def fake_api(method, path, **kwargs):
        calls.append((method, path, kwargs))
        if path == "/api/telegram/control/approvals/77/card":
            return {
                "callbacks": {
                    "approve": "tc1.approve-test",
                    "reject": "tc1.reject-test",
                    "request_changes": "tc1.changes-test",
                }
            }
        assert path == "/api/outreach/campaigns/management-companies/draft"
        return {
            "status": "blocked",
            "task_id": 501,
            "approval_id": 77,
            "recipient_count": 12,
            "subject": "Предложение по клинингу",
            "body": "Добрый день!",
        }

    class File:
        async def download_to_drive(self, custom_path):
            Path(custom_path).write_bytes(b"%PDF-telegram-outreach")

    class Bot:
        async def get_file(self, file_id): return File()

    class Document:
        file_name = "offer.pdf"
        file_size = 100
        file_id = "outreach-file"
        mime_type = "application/pdf"

    class Message:
        document = Document()
        caption = "Разошли это предложение по базе УК"
        replies = []
        async def reply_text(self, value, **kwargs): self.replies.append((value, kwargs))

    class User:
        id = 123

    class Chat:
        id = 123

    class Update:
        effective_message = Message()
        effective_user = User()
        effective_chat = Chat()

    class Context:
        bot = Bot()

    monkeypatch.setattr(settings, "document_storage_path", str(tmp_path))
    monkeypatch.setattr(bot, "allowed", lambda update: True)
    monkeypatch.setattr(bot, "api", fake_api)
    asyncio.run(bot.proposal_document(Update(), Context()))
    assert len(calls) == 2
    assert "12" in Update.effective_message.replies[0][0]
    markup = Update.effective_message.replies[0][1]["reply_markup"]
    assert [button.text for row in markup.inline_keyboard for button in row] == [
        "✅ Одобрить",
        "❌ Отклонить",
        "✏️ Запросить изменения",
    ]


def test_growth_officer_owns_billion_ruble_goal_and_delegates_from_real_contracts(client):
    customer = client.post("/api/entities", json={"entity_type": "client", "name": "Growth customer"}).json()
    site = client.post("/api/entities", json={"entity_type": "site", "name": "Growth site", "parent_id": customer["id"]}).json()
    client.post("/api/entities", json={
        "entity_type": "contract",
        "name": "Growth contract",
        "parent_id": site["id"],
        "status": "active",
        "data": {"monthly_revenue": 1_000_000},
    })
    task = client.post("/api/tasks", json={
        "title": "Growth Officer billion target test",
        "agent_type": "growth_officer",
        "payload": {"action": "billion_revenue_review", "review_at": "2030-01-02T09:00:00"},
        "max_attempts": 1,
    }).json()
    completed = client.post(f"/api/tasks/{task['id']}/run").json()
    assert completed["status"] == "done"
    result = completed["result"]
    assert result["target_rub"] == 1_000_000_000
    assert result["current_rub"] >= 12_000_000
    assert result["source"] == "active contract monthly_revenue fields multiplied by 12"
    assert {row["agent_type"] for row in result["delegated_tasks"]} == {"sales", "marketing", "tender", "finance", "hr"}
    assert result["owner_approval_preserved"] is True
    goal = next(row for row in client.get("/api/goals").json() if row["metric"] == "annual_revenue_run_rate_rub")
    assert goal["owner"] == "growth_officer"
    assert goal["target"] == 1_000_000_000

    report_task = client.post("/api/tasks", json={
        "title": "Growth in 30 minute report test",
        "agent_type": "orchestrator",
        "payload": {"action": "system_activity_report", "period_minutes": 30},
        "max_attempts": 1,
    }).json()
    report = client.post(f"/api/tasks/{report_task['id']}/run").json()["result"]
    assert report["strategic_growth"]["target_rub"] == 1_000_000_000
    assert report["strategic_growth"]["current_rub"] >= 12_000_000
    from app.reports import format_activity_report
    assert "1 млрд" in format_activity_report(report)


def test_worker_skips_rate_limited_mailbox_without_starving_pool(monkeypatch):
    from datetime import datetime, timezone

    from app import worker
    from app.db import SessionLocal
    from app.models import OutboundMessage, SenderMailbox
    from sqlalchemy import select, update

    sent = []

    class SMTP:
        def __init__(self, *args, **kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *args): return False
        def starttls(self): return None
        def login(self, username, password): return None
        def send_message(self, message): sent.append(message["To"])

    monkeypatch.setattr(worker.smtplib, "SMTP", SMTP)
    monkeypatch.setenv("SMTP_LIMITED_TEST", "secret-a")
    monkeypatch.setenv("SMTP_AVAILABLE_TEST", "secret-b")
    now = datetime.now(timezone.utc).replace(tzinfo=None)
    with SessionLocal() as db:
        db.execute(
            update(OutboundMessage)
            .where(OutboundMessage.status.in_(["queued", "waiting_configuration", "retry"]))
            .values(status="sent", sent_at=None)
        )
        limited = SenderMailbox(name="Limited", address="limited@example.com", smtp_host="smtp.example.com", username="limited@example.com", secret_ref="SMTP_LIMITED_TEST", per_minute=100, per_day=1)
        available = SenderMailbox(name="Available", address="available@example.com", smtp_host="smtp.example.com", username="available@example.com", secret_ref="SMTP_AVAILABLE_TEST", per_minute=100, per_day=10)
        db.add_all([limited, available]); db.flush()
        db.add_all([
            OutboundMessage(campaign_key="limit-marker", recipient="already@example.com", subject="Done", body="Done", mailbox_id=limited.id, status="sent", scheduled_at=now, sent_at=now),
            OutboundMessage(campaign_key="limited-queue", recipient="wait@example.com", subject="Wait", body="Wait", mailbox_id=limited.id, status="queued", scheduled_at=now),
            OutboundMessage(campaign_key="available-queue", recipient="send@example.com", subject="Send", body="Send", mailbox_id=available.id, status="queued", scheduled_at=now),
        ])
        db.commit()
        assert worker.send_next_email(db) is True
        assert db.scalar(select(OutboundMessage.status).where(OutboundMessage.campaign_key == "limited-queue")) == "queued"
        assert db.scalar(select(OutboundMessage.status).where(OutboundMessage.campaign_key == "available-queue")) == "sent"
    assert sent == ["send@example.com"]


def test_expired_approval_is_persistently_rejected(client):
    from datetime import datetime, timedelta, timezone

    from sqlalchemy import select

    from app.db import SessionLocal
    from app.models import ApprovalDecisionRecord, ApprovalRequest, Task

    task = client.post("/api/tasks", json={
        "title": "Expired approval test",
        "agent_type": "tender",
        "payload": {"action_kind": "tender_submission"},
    }).json()
    blocked = client.post(f"/api/tasks/{task['id']}/run").json()
    approval_id = blocked["result"]["approval_id"]
    with SessionLocal() as db:
        approval = db.get(ApprovalRequest, approval_id)
        approval.expires_at = datetime.now(timezone.utc).replace(tzinfo=None) - timedelta(seconds=1)
        db.commit()

    response = client.post(
        f"/api/approvals/{approval_id}/approve",
        json={"note": "Too late"},
    )
    assert response.status_code == 409
    assert response.json()["detail"] == "Approval expired"
    with SessionLocal() as db:
        approval = db.get(ApprovalRequest, approval_id)
        decision = db.scalar(
            select(ApprovalDecisionRecord).where(
                ApprovalDecisionRecord.approval_id == approval_id
            )
        )
        assert approval.status == "expired"
        assert approval.decision_version == 2
        assert decision.action == "expire"
        assert decision.actor == "system"
        assert db.get(Task, task["id"]).status == "blocked"


def test_approval_decision_record_and_workflow_resume_are_exactly_once(client):
    from sqlalchemy import func, select

    from app.db import SessionLocal
    from app.models import ApprovalDecisionRecord, TaskTransition

    task = client.post("/api/tasks", json={
        "title": "Exactly once approval test",
        "agent_type": "tender",
        "payload": {"action_kind": "tender_submission"},
    }).json()
    blocked = client.post(f"/api/tasks/{task['id']}/run").json()
    approval_id = blocked["result"]["approval_id"]
    approved = client.post(
        f"/api/approvals/{approval_id}/approve",
        json={"note": "Exact owner decision"},
    )
    assert approved.status_code == 200
    assert approved.json()["decision_version"] == 2
    duplicate = client.post(
        f"/api/approvals/{approval_id}/approve",
        json={"note": "Duplicate click"},
    )
    assert duplicate.status_code == 409

    with SessionLocal() as db:
        assert db.scalar(
            select(func.count()).select_from(ApprovalDecisionRecord).where(
                ApprovalDecisionRecord.approval_id == approval_id
            )
        ) == 1
        assert db.scalar(
            select(func.count()).select_from(TaskTransition).where(
                TaskTransition.transition_key == f"task:{task['id']}:approval:{approval_id}:queued"
            )
        ) == 1


def test_request_changes_records_decision_without_resuming_workflow(client):
    from sqlalchemy import select

    from app.db import SessionLocal
    from app.models import ApprovalDecisionRecord, Task

    task = client.post("/api/tasks", json={
        "title": "Request changes approval test",
        "agent_type": "tender",
        "payload": {"action_kind": "tender_submission"},
    }).json()
    blocked = client.post(f"/api/tasks/{task['id']}/run").json()
    approval_id = blocked["result"]["approval_id"]
    response = client.post(
        f"/api/approvals/{approval_id}/request_changes",
        json={"note": "Уточнить сумму и срок действия"},
    )
    assert response.status_code == 200
    assert response.json()["status"] == "changes_requested"
    assert response.json()["execution"] == "not_executed"
    with SessionLocal() as db:
        decision = db.scalar(
            select(ApprovalDecisionRecord).where(
                ApprovalDecisionRecord.approval_id == approval_id
            )
        )
        assert decision.action == "request_changes"
        assert decision.reason == "Уточнить сумму и срок действия"
        assert db.get(Task, task["id"]).status == "blocked"


def test_non_owner_role_cannot_decide_protected_approval(client):
    task = client.post("/api/tasks", json={
        "title": "Manager approval denial test",
        "agent_type": "tender",
        "payload": {"action_kind": "tender_submission"},
    }).json()
    blocked = client.post(f"/api/tasks/{task['id']}/run").json()
    response = client.post(
        f"/api/approvals/{blocked['result']['approval_id']}/approve",
        headers={"X-Role": "manager"},
        json={"note": "Manager must not approve"},
    )
    assert response.status_code == 403
